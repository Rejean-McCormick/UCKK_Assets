#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
king_klown_premium_builder.py

Moteur de standardisation premium pour manuels King Klown / UCKK.

Entrée :
    .docx exporté de Google Docs ou Word

Sorties :
    .docx standardisé
    .pdf optionnel si LibreOffice est installé

Dépendance :
    pip install python-docx

Lancer le GUI :
    python king_klown_premium_builder.py

CLI :
    python king_klown_premium_builder.py "manuel.docx" --toc
    python king_klown_premium_builder.py "./docs" --recursive --margin-mode recto-verso --pdf

Notes :
- Le script garde seulement le layout Premium.
- Texte forcé en noir; aucun texte en couleur de marque.
- Il corrige les paragraphes courts justifiés qui créent des espaces énormes.
- Par défaut, le corps est aligné à gauche pour éviter les blancs étirés.
- Il ignore les titres génériques comme "Word Document" dans le header.
- Le header peut suivre automatiquement le chapitre courant via STYLEREF.
- Il utilise par défaut le layout dominant du modèle fourni : KDP 6 x 9 po, marges miroir, en-tête/pied normalisés.
- Les H1 de chapitre utilisent un Section Break (Next Page), pas un faux retour de ligne.
- Les H2 optionnels utilisent un Page Break manuel supprimable, pas pageBreakBefore dans le style.
"""

from __future__ import annotations

import argparse
import os
import queue
import re
import shutil
import subprocess
import sys
import threading
import time
import tkinter as tk
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
from typing import Callable, Iterable, Optional
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from docx import Document
    from docx.document import Document as DocxDocument
    from docx.enum.section import WD_SECTION_START
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor, Twips
except ImportError as exc:
    raise SystemExit(
        "Dépendance manquante : python-docx.\n"
        "Installe-la avec : pip install python-docx"
    ) from exc


# ---------------------------------------------------------------------------
# Charte Premium King Klown / UCKK
# ---------------------------------------------------------------------------

UCKK_DARK = "1B1B1B"
UCKK_SOFT_BLACK = "2A2A2A"
UCKK_GRAY = "666666"
UCKK_LIGHT_GRAY = "F4F4F2"
UCKK_BORDER = "D7D2C8"
UCKK_GOLD = "1E6864"  # accent principal
UCKK_DEEP_RED = "1E6864"  # couleur de marque King Klown / UCKK
UCKK_PARCHMENT = "E8F3F1"
UCKK_PALE_GOLD = "E8F3F1"
WHITE = "FFFFFF"
BLACK = "000000"

# ---------------------------------------------------------------------------
# PARAMÈTRES PAR DÉFAUT — À MODIFIER ICI
# ---------------------------------------------------------------------------

DEFAULT_OUTPUT_DIR = Path.home() / "Downloads"
DEFAULT_SUFFIX = "UCKK_formated"

DEFAULT_PAGE_SIZE = "kdp_6x9"
DEFAULT_MARGIN_MODE = "recto-verso"
DEFAULT_HEADER_LAYOUT = "chapter"
DEFAULT_BODY_ALIGNMENT = "left"

DEFAULT_MARGIN_TOP_CM = 1.40
DEFAULT_MARGIN_BOTTOM_CM = 1.40
DEFAULT_MARGIN_LEFT_CM = 1.52
DEFAULT_MARGIN_RIGHT_CM = 1.25
DEFAULT_MARGIN_INSIDE_CM = 1.52
DEFAULT_MARGIN_OUTSIDE_CM = 1.25
DEFAULT_HEADER_DISTANCE_CM = 0.64
DEFAULT_FOOTER_DISTANCE_CM = 0.64

DEFAULT_BODY_FONT = "EB Garamond"
DEFAULT_HEADING_FONT = "EB Garamond"
DEFAULT_MONO_FONT = "Consolas"
DEFAULT_FOOTER_TEXT = "Univers-Cité King Klown"

DEFAULT_BODY_SIZE_PT = 10.7
DEFAULT_LINE_SPACING = 1.045
DEFAULT_BODY_SPACE_AFTER_PT = 3.5

DEFAULT_H1_SIZE_PT = 17.5
DEFAULT_H1_SPACE_BEFORE_PT = 24.0
DEFAULT_H1_SPACE_AFTER_PT = 9.0
DEFAULT_H1_TOP_GAP_PT = 24.0

DEFAULT_H2_SIZE_PT = 13.2
DEFAULT_H2_SPACE_BEFORE_PT = 16.0
DEFAULT_H2_SPACE_AFTER_PT = 5.0
DEFAULT_H2_PAGE_BREAKS = False

DEFAULT_H3_SIZE_PT = 11.6
DEFAULT_H3_SPACE_BEFORE_PT = 12.0
DEFAULT_H3_SPACE_AFTER_PT = 3.5

DEFAULT_H4_SIZE_PT = 10.6
DEFAULT_H4_SPACE_BEFORE_PT = 9.0
DEFAULT_H4_SPACE_AFTER_PT = 2.0

DEFAULT_LIST_BLOCK_SPACE_AFTER_PT = 4.0

DEFAULT_ADD_TOC_GUI = True
DEFAULT_ADD_HEADER_FOOTER = True
DEFAULT_CHAPTER_PAGE_BREAKS = True
DEFAULT_DO_NOT_CHANGE_HEADING_LEVELS = True
DEFAULT_CLEAR_GOOGLE_FORMATTING = True
DEFAULT_STANDARDIZE_TABLES = True
DEFAULT_CALLOUT_BOXES = False
DEFAULT_EXPORT_PDF = False
DEFAULT_DRY_RUN = False
DEFAULT_RECURSIVE = False
DEFAULT_IN_PLACE = False


# Layout dominant observé dans le document modèle fourni.
# Les valeurs exactes en twips évitent les petits écarts de conversion.
KDP_6X9_WIDTH_TWIPS = 8640
KDP_6X9_HEIGHT_TWIPS = 12960
KDP_6X9_TOP_TWIPS = 794
KDP_6X9_BOTTOM_TWIPS = 794
KDP_6X9_INSIDE_TWIPS = 862
KDP_6X9_OUTSIDE_TWIPS = 709
KDP_6X9_HEADER_TWIPS = 284
KDP_6X9_FOOTER_TWIPS = 227
KDP_6X9_GUTTER_TWIPS = 0


@dataclass
class PremiumConfig:
    page_size: str = DEFAULT_PAGE_SIZE            # kdp_6x9, letter ou a4
    margin_mode: str = DEFAULT_MARGIN_MODE      # recto-verso, centered, left

    # Layout dominant du modèle fourni : KDP 6 x 9 po, marges miroir.
    # Valeurs dominantes repérées dans les chapitres :
    # page 8640 x 12960 twips ; marges 794 / 794 / 862 / 709 twips ;
    # en-tête 284 twips ; pied 227 twips ; gouttière 0.
    margin_top_cm: float = DEFAULT_MARGIN_TOP_CM
    margin_bottom_cm: float = DEFAULT_MARGIN_BOTTOM_CM
    margin_left_cm: float = DEFAULT_MARGIN_LEFT_CM
    margin_right_cm: float = DEFAULT_MARGIN_RIGHT_CM
    margin_inside_cm: float = DEFAULT_MARGIN_INSIDE_CM
    margin_outside_cm: float = DEFAULT_MARGIN_OUTSIDE_CM
    header_distance_cm: float = DEFAULT_HEADER_DISTANCE_CM
    footer_distance_cm: float = DEFAULT_FOOTER_DISTANCE_CM

    body_font: str = DEFAULT_BODY_FONT
    heading_font: str = DEFAULT_HEADING_FONT
    mono_font: str = DEFAULT_MONO_FONT

    body_size_pt: float = DEFAULT_BODY_SIZE_PT
    line_spacing: float = DEFAULT_LINE_SPACING
    body_space_after_pt: float = DEFAULT_BODY_SPACE_AFTER_PT

    add_toc: bool = False
    add_header_footer: bool = DEFAULT_ADD_HEADER_FOOTER
    header_layout: str = DEFAULT_HEADER_LAYOUT         # chapter, brand-chapter, centered, split, minimal, none
    body_alignment: str = DEFAULT_BODY_ALIGNMENT          # left, smart-justify, justify
    footer_text: str = DEFAULT_FOOTER_TEXT

    chapter_page_breaks: bool = DEFAULT_CHAPTER_PAGE_BREAKS
    h1_new_page_top_gap_pt: float = DEFAULT_H1_TOP_GAP_PT
    page_breaks_before_h2: bool = DEFAULT_H2_PAGE_BREAKS
    list_block_space_after_pt: float = DEFAULT_LIST_BLOCK_SPACE_AFTER_PT
    clear_google_run_formatting: bool = DEFAULT_CLEAR_GOOGLE_FORMATTING
    normalize_heading_levels: bool = False
    standardize_tables: bool = DEFAULT_STANDARDIZE_TABLES
    add_callout_boxes: bool = DEFAULT_CALLOUT_BOXES
    export_pdf: bool = DEFAULT_EXPORT_PDF


@dataclass
class ProcessStats:
    source: Path
    output: Path
    headings: int = 0
    callouts: int = 0
    formula_blocks: int = 0
    code_lines: int = 0
    tables: int = 0
    blank_paragraphs_removed: int = 0
    warnings: list[str] = field(default_factory=list)
    dry_run: bool = False


# ---------------------------------------------------------------------------
# OOXML helpers
# ---------------------------------------------------------------------------

def _hex(rgb: str) -> str:
    return rgb.replace("#", "").upper()


def get_or_add(parent, child_tag: str):
    child = parent.find(qn(child_tag))
    if child is None:
        child = OxmlElement(child_tag)
        parent.append(child)
    return child


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = get_or_add(tc_pr, "w:shd")
    shd.set(qn("w:fill"), _hex(fill))


def set_cell_borders(cell, color: str = UCKK_BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = get_or_add(tc_pr, "w:tcBorders")

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), _hex(color))


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = get_or_add(tc_pr, "w:tcMar")
    values = {"top": top, "start": start, "bottom": bottom, "end": end}

    for side, value in values.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = get_or_add(p_pr, "w:shd")
    shd.set(qn("w:fill"), _hex(fill))


def set_paragraph_border(paragraph, side: str = "left", color: str = UCKK_GOLD, size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = get_or_add(p_pr, "w:pBdr")

    border = borders.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        borders.append(border)

    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "6")
    border.set(qn("w:color"), _hex(color))


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def add_field(paragraph, instr: str, placeholder: str = "") -> None:
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    run._r.append(instr_text)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        run._r.append(t)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def enable_update_fields_on_open(docx_path: Path) -> None:
    tmp_path = docx_path.with_suffix(".tmp.docx")
    with ZipFile(docx_path, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                xml = data.decode("utf-8")
                if "w:updateFields" not in xml:
                    xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
                    data = xml.encode("utf-8")
            zout.writestr(item, data)
    tmp_path.replace(docx_path)


def enable_mirror_margins(doc: DocxDocument) -> None:
    settings = doc.settings._element
    if settings.find(qn("w:mirrorMargins")) is None:
        settings.append(OxmlElement("w:mirrorMargins"))


def clear_run_formatting_keep_emphasis(run) -> None:
    r_pr = run._r.rPr
    if r_pr is None:
        return

    # On supprime les surcharges qui viennent souvent de Google Docs,
    # mais on garde gras, italique, souligné, exposant, indice.
    for tag in [
        "w:rFonts", "w:sz", "w:szCs", "w:color", "w:highlight", "w:shd",
        "w:spacing", "w:kern", "w:position", "w:lang",
    ]:
        for child in list(r_pr.findall(qn(tag))):
            r_pr.remove(child)


# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

def style_lookup(doc: DocxDocument, *names: str):
    """Trouve un style par son nom visible, sans passer par style_id.

    Les versions récentes de python-docx émettent un avertissement quand
    doc.styles[...] tombe sur un style_id. On parcourt donc explicitement
    les styles et on compare uniquement leur nom visible.
    """
    lower = {s.name.lower(): s for s in doc.styles if getattr(s, "name", None)}
    for name in names:
        if not name:
            continue
        found = lower.get(name.lower())
        if found is not None:
            return found
    return None


def ensure_paragraph_style(doc: DocxDocument, name: str, base: Optional[str] = None):
    style = style_lookup(doc, name)
    if style is not None:
        return style

    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base:
        base_style = style_lookup(doc, base)
        if base_style is not None:
            style.base_style = base_style
    return style


def set_font(style, name: str, size_pt: Optional[float] = None,
             bold: Optional[bool] = None, italic: Optional[bool] = None,
             color: Optional[str] = None) -> None:
    style.font.name = name
    if style._element.rPr is not None and style._element.rPr.rFonts is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        style._element.rPr.rFonts.set(qn("w:cs"), name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(_hex(color))


def remove_style_page_break_before(style) -> None:
    """Retire pageBreakBefore du style lui-même.

    Les sauts de page optionnels doivent rester des objets supprimables dans le
    document, pas une propriété forcée du style Heading/Titre.
    """
    if style is None:
        return
    p_pr = getattr(style._element, "pPr", None)
    if p_pr is None:
        return
    for element in list(p_pr.findall(qn("w:pageBreakBefore"))):
        p_pr.remove(element)


def remove_heading_style_page_breaks(doc: DocxDocument) -> None:
    """Nettoie les pageBreakBefore hérités des styles Heading/Titre."""
    for level in range(1, 7):
        remove_style_page_break_before(style_lookup(doc, f"Heading {level}", f"Titre {level}", f"H{level}"))
    remove_style_page_break_before(style_lookup(doc, "Subtitle", "Sous-titre"))


def set_ooxml_run_font(r_pr, name: str, size_pt: Optional[float] = None, color: str = BLACK) -> None:
    """Force police/couleur dans un élément w:rPr OOXML.

    Utile pour les puces et les listes numérotées : dans Word, le numéro ou
    la puce peut prendre sa police depuis numbering.xml, pas depuis le texte
    du paragraphe.
    """
    r_fonts = get_or_add(r_pr, "w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), name)

    if size_pt is not None:
        half_points = str(int(round(size_pt * 2)))
        sz = get_or_add(r_pr, "w:sz")
        sz.set(qn("w:val"), half_points)
        sz_cs = get_or_add(r_pr, "w:szCs")
        sz_cs.set(qn("w:val"), half_points)

    color_el = get_or_add(r_pr, "w:color")
    color_el.set(qn("w:val"), _hex(color))


def force_paragraph_runs_font(paragraph, cfg: PremiumConfig, size_pt: Optional[float] = None) -> None:
    """Force la police du texte d'un paragraphe sans changer sa structure."""
    target_size = cfg.body_size_pt if size_pt is None else size_pt
    for run in paragraph.runs:
        run.font.name = cfg.body_font
        run.font.size = Pt(target_size)
        run.font.color.rgb = RGBColor.from_string(BLACK)
        if run._r.rPr is not None:
            set_ooxml_run_font(run._r.rPr, cfg.body_font, target_size, BLACK)


def force_paragraph_mark_font(paragraph, cfg: PremiumConfig, size_pt: Optional[float] = None) -> None:
    """Force aussi la police du marqueur de paragraphe.

    Word utilise souvent les propriétés du marqueur de paragraphe pour
    dessiner les numéros ou les puces. Sans ça, le texte peut être dans la
    bonne police, mais le numéro peut rester en Calibri/Aptos.
    """
    target_size = cfg.body_size_pt if size_pt is None else size_pt
    p_pr = paragraph._p.get_or_add_pPr()
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)
    set_ooxml_run_font(r_pr, cfg.body_font, target_size, BLACK)


def force_list_paragraph_font(paragraph, cfg: PremiumConfig, number_format: Optional[str] = None) -> None:
    """Uniformise le texte des listes et les marqueurs numérotés.

    Pour les bullets, on laisse le marqueur venir de numbering.xml afin de
    pouvoir lui donner une taille de symbole plus discrète. Pour les listes
    numérotées, on force aussi le marqueur de paragraphe, sinon les numéros
    peuvent rester en Calibri/Aptos.
    """
    force_paragraph_runs_font(paragraph, cfg)
    if number_format != "bullet":
        force_paragraph_mark_font(paragraph, cfg)


def get_numbering_level_format(paragraph) -> Optional[str]:
    """Retourne le format du niveau de liste : bullet, decimal, lowerLetter, etc."""
    p_pr = paragraph._p.pPr
    if p_pr is None or p_pr.numPr is None:
        return None

    num_id_el = p_pr.numPr.numId
    if num_id_el is None:
        return None
    num_id = num_id_el.val

    ilvl_el = p_pr.numPr.ilvl
    ilvl = ilvl_el.val if ilvl_el is not None else 0

    try:
        numbering = paragraph.part.numbering_part.element
    except Exception:
        return None

    abstract_id = None
    for num in numbering.findall(qn('w:num')):
        if num.get(qn('w:numId')) == str(num_id):
            abstract_id_el = num.find(qn('w:abstractNumId'))
            if abstract_id_el is not None:
                abstract_id = abstract_id_el.get(qn('w:val'))
            break
    if abstract_id is None:
        return None

    for abstract in numbering.findall(qn('w:abstractNum')):
        if abstract.get(qn('w:abstractNumId')) != str(abstract_id):
            continue
        for lvl in abstract.findall(qn('w:lvl')):
            if lvl.get(qn('w:ilvl')) == str(ilvl):
                num_fmt = lvl.find(qn('w:numFmt'))
                return num_fmt.get(qn('w:val')) if num_fmt is not None else None
    return None


def standardize_numbering_fonts(doc: DocxDocument, cfg: PremiumConfig) -> None:
    """Force la police des puces et numéros dans numbering.xml.

    Corrige le cas où les numbered lists affichent un numéro dans une autre
    police que le texte, même si les runs du paragraphe sont bien formatés.
    """
    try:
        numbering = doc.part.numbering_part.element
    except Exception:
        return

    for lvl in numbering.findall('.//' + qn('w:lvl')):
        r_pr = lvl.find(qn('w:rPr'))
        if r_pr is None:
            r_pr = OxmlElement('w:rPr')
            lvl.append(r_pr)
        num_fmt = lvl.find(qn('w:numFmt'))
        fmt = num_fmt.get(qn('w:val')) if num_fmt is not None else None

        if fmt == 'bullet':
            # Les bullets en EB Garamond deviennent souvent trop grosses.
            # On garde le texte en EB Garamond, mais on dessine le symbole
            # avec une police de symbole plus neutre et une taille réduite.
            set_ooxml_run_font(r_pr, 'Arial', 8.0, BLACK)
            lvl_text = lvl.find(qn('w:lvlText'))
            if lvl_text is not None and lvl_text.get(qn('w:val')) in {'●', '•'}:
                lvl_text.set(qn('w:val'), '•')
        else:
            # Les numéros doivent suivre la police du corps.
            set_ooxml_run_font(r_pr, cfg.body_font, cfg.body_size_pt, BLACK)

        # Nettoie les polices thématiques qui peuvent reprendre Aptos/Calibri
        # malgré w:ascii/w:hAnsi.
        r_fonts = r_pr.find(qn('w:rFonts'))
        if r_fonts is not None:
            for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme', 'w:csTheme'):
                key = qn(attr)
                if key in r_fonts.attrib:
                    del r_fonts.attrib[key]


def alignment_from_body_setting(body_alignment: str):
    """Retourne l'alignement Word pour le corps.

    left est le défaut premium, parce que les exports Google Docs/Word
    contiennent souvent des retours de ligne manuels. En justification,
    Word étire alors chaque ligne courte et crée des blancs énormes.
    """
    if body_alignment == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if body_alignment == "smart-justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def configure_styles(doc: DocxDocument, cfg: PremiumConfig) -> None:
    remove_heading_style_page_breaks(doc)
    normal = style_lookup(doc, "Normal", "normal") or ensure_paragraph_style(doc, "Normal")
    set_font(normal, cfg.body_font, cfg.body_size_pt, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(cfg.body_space_after_pt)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = cfg.line_spacing
    normal.paragraph_format.alignment = alignment_from_body_setting(cfg.body_alignment)

    h1 = style_lookup(doc, "Heading 1", "Titre 1") or ensure_paragraph_style(doc, "Heading 1", "Normal")
    set_font(h1, cfg.heading_font, DEFAULT_H1_SIZE_PT, bold=True, color=BLACK)
    h1.paragraph_format.space_before = Pt(DEFAULT_H1_SPACE_BEFORE_PT)
    h1.paragraph_format.space_after = Pt(DEFAULT_H1_SPACE_AFTER_PT)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h2 = style_lookup(doc, "Heading 2", "Titre 2") or ensure_paragraph_style(doc, "Heading 2", "Normal")
    set_font(h2, cfg.heading_font, DEFAULT_H2_SIZE_PT, bold=True, color=BLACK)
    h2.paragraph_format.space_before = Pt(DEFAULT_H2_SPACE_BEFORE_PT)
    h2.paragraph_format.space_after = Pt(DEFAULT_H2_SPACE_AFTER_PT)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h3 = style_lookup(doc, "Heading 3", "Titre 3") or ensure_paragraph_style(doc, "Heading 3", "Normal")
    set_font(h3, cfg.heading_font, DEFAULT_H3_SIZE_PT, bold=True, color=BLACK)
    h3.paragraph_format.space_before = Pt(DEFAULT_H3_SPACE_BEFORE_PT)
    h3.paragraph_format.space_after = Pt(DEFAULT_H3_SPACE_AFTER_PT)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h4 = style_lookup(doc, "Heading 4", "Titre 4") or ensure_paragraph_style(doc, "Heading 4", "Normal")
    set_font(h4, cfg.heading_font, DEFAULT_H4_SIZE_PT, bold=True, color=BLACK)
    h4.paragraph_format.space_before = Pt(DEFAULT_H4_SPACE_BEFORE_PT)
    h4.paragraph_format.space_after = Pt(DEFAULT_H4_SPACE_AFTER_PT)
    h4.paragraph_format.keep_with_next = True
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = ensure_paragraph_style(doc, "KK - Intro", "Normal")
    set_font(intro, cfg.body_font, cfg.body_size_pt + 0.5, color=BLACK)
    intro.paragraph_format.space_before = Pt(4)
    intro.paragraph_format.space_after = Pt(9)
    intro.paragraph_format.line_spacing = 1.08
    intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    callout = ensure_paragraph_style(doc, "KK - Encadre", "Normal")
    set_font(callout, cfg.body_font, cfg.body_size_pt, color=BLACK)
    callout.paragraph_format.left_indent = Cm(0.38)
    callout.paragraph_format.right_indent = Cm(0.25)
    callout.paragraph_format.space_before = Pt(7)
    callout.paragraph_format.space_after = Pt(7)
    callout.paragraph_format.line_spacing = 1.04
    callout.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    formula = ensure_paragraph_style(doc, "KK - Formule", "Normal")
    set_font(formula, cfg.body_font, cfg.body_size_pt + 0.2, italic=False, color=BLACK)
    formula.paragraph_format.left_indent = Cm(0.4)
    formula.paragraph_format.space_before = Pt(5)
    formula.paragraph_format.space_after = Pt(5)
    formula.paragraph_format.line_spacing = 1.05
    formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    transformation = ensure_paragraph_style(doc, "KK - Transformation", "Normal")
    set_font(transformation, cfg.body_font, cfg.body_size_pt, italic=True, color=BLACK)
    # Les blocs avec flèches ressemblent souvent à des listes/trajectoires.
    # Ils doivent rester alignés à gauche, pas centrés.
    transformation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    transformation.paragraph_format.space_before = Pt(6)
    transformation.paragraph_format.space_after = Pt(6)
    transformation.paragraph_format.line_spacing = 1.0

    code = ensure_paragraph_style(doc, "KK - Code", "Normal")
    set_font(code, cfg.mono_font, 9.2, color=BLACK)
    code.paragraph_format.left_indent = Cm(0.35)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    caption = ensure_paragraph_style(doc, "KK - Legende", "Normal")
    set_font(caption, cfg.body_font, 9.6, italic=True, color=BLACK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)

    # Listes compactes : évite le grand blanc avant/entre les puces.
    # Certains documents Word/Google Docs gardent un espace hérité dans
    # les styles de listes, même si les paragraphes sont ensuite traités.
    for list_style_name in (
        "List Paragraph", "Liste à puces", "Liste numérotée",
        "List Bullet", "List Bullet 2", "List Bullet 3",
        "List Number", "List Number 2", "List Number 3",
    ):
        list_style = style_lookup(doc, list_style_name)
        if list_style is None:
            continue
        set_font(list_style, cfg.body_font, cfg.body_size_pt, color=BLACK)
        list_style.paragraph_format.space_before = Pt(0)
        list_style.paragraph_format.space_after = Pt(0)
        list_style.paragraph_format.line_spacing = 1.0
        list_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ---------------------------------------------------------------------------
# Détection de structure
# ---------------------------------------------------------------------------

H4_RE = re.compile(r"^\s*\d{1,2}\.\d+\.\d+\.\d+\s+.+")
H3_RE = re.compile(r"^\s*\d{1,2}\.\d+\.\d+\s+.+")
H2_RE = re.compile(r"^\s*\d{1,2}\.\d+\s+.+")
H1_RE = re.compile(r"^\s*\d{2}\.\s+.+")
PART_RE = re.compile(r"^\s*Partie\s+\d+\s+[—-]\s+.+", re.IGNORECASE)
CHAPTER_RE = re.compile(r"^\s*(Chapitre|Chapter)\s+\d+\s+[—-]\s+.+", re.IGNORECASE)

CALLOUT_RE = re.compile(
    r"^\s*(Phrase clé|Point de vigilance|Formule courte|Formule officielle|"
    r"Devise|Règle|À retenir|Retenir|Erreur fréquente|Erreurs fréquentes|"
    r"Production attendue|Critères de réussite|Critères|Trace à conserver|"
    r"Livrable|Mini-exercice|Exercice|Objectifs du chapitre|Attention|Vigilance)\b",
    re.IGNORECASE,
)

GENERIC_TITLES = {
    "word document",
    "document",
    "untitled document",
    "sans titre",
    "document sans titre",
}


def clean_markdown_artifacts(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = re.sub(r"^\*\*(.+)\*\*$", r"\1", text)
    text = re.sub(r"^__(.+)__$", r"\1", text)
    text = text.replace("\\#", "#")
    text = text.replace("\\_", "_")
    text = text.replace("\\*", "*")
    return text.strip()


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        for i, run in enumerate(paragraph.runs):
            run.text = new_text if i == 0 else ""
    else:
        paragraph.add_run(new_text)


def strip_markdown_heading_marks(paragraph) -> None:
    cleaned = clean_markdown_artifacts(paragraph.text)
    replace_paragraph_text(paragraph, cleaned)


def classify_heading(text: str) -> Optional[str]:
    stripped = text.strip()
    if not stripped:
        return None

    if stripped.startswith("#### "):
        return "Heading 4"
    if stripped.startswith("### "):
        return "Heading 3"
    if stripped.startswith("## "):
        return "Heading 2"
    if stripped.startswith("# "):
        return "Heading 1"

    cleaned = clean_markdown_artifacts(stripped)

    if H4_RE.match(cleaned):
        return "Heading 4"
    if H3_RE.match(cleaned):
        return "Heading 3"
    if H2_RE.match(cleaned):
        return "Heading 2"
    if H1_RE.match(cleaned) or PART_RE.match(cleaned) or CHAPTER_RE.match(cleaned):
        return "Heading 1"
    return None


def is_list_paragraph(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.numPr is not None


def has_manual_line_breaks(paragraph) -> bool:
    """Google Docs et Word exportent parfois des retours de ligne manuels.
    En texte justifié, ces lignes sont étirées. On les repère pour les mettre
    à gauche ou en formule.
    """
    if "\n" in paragraph.text:
        return True
    return any("\n" in run.text for run in paragraph.runs)


def is_spread_risk_text(text: str) -> bool:
    cleaned = clean_markdown_artifacts(text)
    if not cleaned:
        return False

    lines = [ln.strip() for ln in cleaned.splitlines() if ln.strip()]
    if len(lines) > 1:
        # Plusieurs lignes courtes dans un même paragraphe = grand risque
        # d'espaces étirés en justification.
        return any(len(ln) <= 120 for ln in lines)

    if len(cleaned) <= 140:
        return True

    # Les phrases avec peu de mots longs et beaucoup d'espaces potentiels
    # se déforment facilement dans un bloc justifié.
    words = cleaned.split()
    return len(words) <= 12



def looks_like_formula_block(text: str) -> bool:
    """Détecte les formules courtes qui ne doivent jamais être justifiées.

    Corrige notamment le problème visible dans Word :
    "Inventorier      avant      d'analyser"
    causé par un paragraphe justifié avec retours de ligne manuels.
    """
    cleaned = clean_markdown_artifacts(text)
    lines = [ln.strip() for ln in cleaned.splitlines() if ln.strip()]

    if not lines:
        return False

    if "→" in cleaned and len(cleaned) <= 220:
        return True

    if len(lines) >= 2 and all(len(ln) <= 90 for ln in lines):
        formula_hits = sum(
            1 for ln in lines
            if re.search(r"\bavant\b|\b→\b|:$", ln, flags=re.IGNORECASE)
        )
        return formula_hits >= 1

    if len(cleaned) <= 110 and re.search(r"\bavant\b", cleaned, flags=re.IGNORECASE):
        return True

    return False


def looks_like_code_line(text: str) -> bool:
    cleaned = clean_markdown_artifacts(text)
    return bool(re.match(r"^\s*(```|[a-zA-Z0-9_]+\s*=\s*|[A-Za-z0-9_./-]+\.(py|md|docx|pdf))", cleaned))


def previous_non_empty_paragraph(paragraphs: list, index: int):
    """Retourne le paragraphe non vide précédent, utilisé pour resserrer
    l'espace visible avant le premier item d'une liste.
    """
    for prev in reversed(paragraphs[:index]):
        if prev.text.strip():
            return prev
    return None


def previous_paragraph_is_heading(paragraph) -> bool:
    style_name = paragraph_style_name(paragraph)
    style_id = paragraph_style_id(paragraph)
    return (
        style_name.startswith("heading")
        or style_name.startswith("titre")
        or style_id.startswith("heading")
        or style_id.startswith("titre")
    )


def paragraph_style_name(paragraph) -> str:
    try:
        return (paragraph.style.name or "").strip().lower()
    except Exception:
        return ""


def paragraph_style_id(paragraph) -> str:
    try:
        return (paragraph.style.style_id or "").strip().lower()
    except Exception:
        return ""


def style_or_base_is_heading_1(style) -> bool:
    """Reconnaît Heading 1 / Titre 1, même via style_id ou style basé dessus."""
    seen: set[int] = set()
    while style is not None:
        ident = id(style)
        if ident in seen:
            return False
        seen.add(ident)

        name = (getattr(style, "name", "") or "").strip().lower()
        style_id = (getattr(style, "style_id", "") or "").strip().lower()
        if name in {"heading 1", "titre 1", "h1"}:
            return True
        if style_id in {"heading1", "titre1", "h1"}:
            return True

        style = getattr(style, "base_style", None)
    return False


def is_heading_1_paragraph(paragraph) -> bool:
    try:
        return style_or_base_is_heading_1(paragraph.style)
    except Exception:
        return paragraph_style_name(paragraph) in {"heading 1", "titre 1", "h1"}


def heading_level(paragraph) -> Optional[int]:
    """Retourne 1-6 pour les vrais styles Word Heading/Titre, sinon None."""
    try:
        style = paragraph.style
    except Exception:
        style = None

    while style is not None:
        name = (getattr(style, "name", "") or "").strip().lower()
        style_id = (getattr(style, "style_id", "") or "").strip().lower()
        for level in range(1, 7):
            if name in {f"heading {level}", f"titre {level}", f"h{level}"}:
                return level
            if style_id in {f"heading{level}", f"titre{level}", f"h{level}"}:
                return level
        style = getattr(style, "base_style", None)
    return None


def apply_heading_paragraph_format(paragraph, cfg: PremiumConfig, level: int) -> None:
    """Applique les espacements des titres sans ajouter de paragraphe vide."""
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.keep_with_next = True

    if level == 1:
        paragraph.paragraph_format.space_before = Pt(DEFAULT_H1_SPACE_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(DEFAULT_H1_SPACE_AFTER_PT)
        for run in paragraph.runs:
            run.font.name = cfg.heading_font
            run.font.size = Pt(DEFAULT_H1_SIZE_PT)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(BLACK)
            if run._r.rPr is not None:
                set_ooxml_run_font(run._r.rPr, cfg.heading_font, DEFAULT_H1_SIZE_PT, BLACK)
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(DEFAULT_H2_SPACE_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(DEFAULT_H2_SPACE_AFTER_PT)
        # Important : on ne force plus le saut de page dans le format du H2.
        # Le mode optionnel H2 utilise un vrai saut de page explicite juste
        # avant le titre, donc l'utilisateur peut le supprimer au cas par cas.
        remove_direct_page_break_before(paragraph)
        for run in paragraph.runs:
            run.font.name = cfg.heading_font
            run.font.size = Pt(DEFAULT_H2_SIZE_PT)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(BLACK)
            if run._r.rPr is not None:
                set_ooxml_run_font(run._r.rPr, cfg.heading_font, DEFAULT_H2_SIZE_PT, BLACK)
    elif level == 3:
        paragraph.paragraph_format.space_before = Pt(DEFAULT_H3_SPACE_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(DEFAULT_H3_SPACE_AFTER_PT)
        for run in paragraph.runs:
            run.font.name = cfg.heading_font
            run.font.size = Pt(DEFAULT_H3_SIZE_PT)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(BLACK)
            if run._r.rPr is not None:
                set_ooxml_run_font(run._r.rPr, cfg.heading_font, DEFAULT_H3_SIZE_PT, BLACK)
    elif level == 4:
        paragraph.paragraph_format.space_before = Pt(DEFAULT_H4_SPACE_BEFORE_PT)
        paragraph.paragraph_format.space_after = Pt(DEFAULT_H4_SPACE_AFTER_PT)
        for run in paragraph.runs:
            run.font.name = cfg.heading_font
            run.font.size = Pt(DEFAULT_H4_SIZE_PT)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(BLACK)
            if run._r.rPr is not None:
                set_ooxml_run_font(run._r.rPr, cfg.heading_font, DEFAULT_H4_SIZE_PT, BLACK)


def paragraph_text_from_element(p_elm) -> str:
    return "".join(t.text or "" for t in p_elm.iter(qn("w:t")))


def paragraph_has_page_break_element(p_elm) -> bool:
    p_pr = p_elm.find(qn("w:pPr"))
    if p_pr is not None:
        page_break_before = p_pr.find(qn("w:pageBreakBefore"))
        if page_break_before is not None:
            val = (page_break_before.get(qn("w:val")) or "1").lower()
            if val not in {"0", "false", "off"}:
                return True

    for br in p_elm.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def paragraph_has_explicit_page_break_run(p_elm) -> bool:
    """Vrai seulement pour un saut de page réel dans un run, pas pageBreakBefore."""
    if p_elm is None:
        return False
    for br in p_elm.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def paragraph_ends_with_explicit_page_break_run(p_elm) -> bool:
    """Retourne True si le paragraphe contient déjà un saut de page manuel."""
    return paragraph_has_explicit_page_break_run(p_elm)


def paragraph_starts_with_explicit_page_break_run(p_elm) -> bool:
    """Repère un saut de page manuel placé au début du paragraphe courant."""
    if p_elm is None:
        return False
    for child in p_elm:
        if child.tag == qn("w:pPr"):
            continue
        if child.tag == qn("w:r"):
            br = child.find(qn("w:br"))
            if br is not None and br.get(qn("w:type")) == "page":
                return True
            # Si le premier run contient du texte, on arrête : le break n'est pas au début.
            if child.find(qn("w:t")) is not None:
                return False
        elif child.tag == qn("w:bookmarkStart") or child.tag == qn("w:bookmarkEnd"):
            continue
        else:
            return False
    return False


def insert_manual_page_break_before(paragraph) -> bool:
    """Insère un vrai Page Break avant le paragraphe, sans toucher au style.

    Le saut est placé à la fin du paragraphe précédent, comme un page break
    manuel ordinaire. Dans Word, l'utilisateur peut donc le sélectionner et le
    supprimer pour faire une exception à un H2. Aucun paragraphe vide n'est créé.
    """
    previous = paragraph._p.getprevious()
    if previous is None or previous.tag != qn("w:p"):
        return False

    # Ne pas empiler les sauts si le script est relancé.
    if paragraph_ends_with_explicit_page_break_run(previous):
        return False
    if paragraph_starts_with_explicit_page_break_run(paragraph._p):
        return False
    if paragraph_has_section_break(previous):
        return False

    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    previous.append(run)
    return True


def is_h1_top_gap_spacer_element(p_elm) -> bool:
    """Repère les paragraphes-espace insérés avant les H1 par ce script."""
    if p_elm is None or p_elm.tag != qn("w:p"):
        return False
    text = paragraph_text_from_element(p_elm).replace("\u00A0", "").replace("\u200B", "").strip()
    return text == "" and paragraph_has_page_break_element(p_elm)


def remove_existing_h1_page_spacer_before(paragraph) -> None:
    """Évite d'empiler des espaces si le script est relancé plusieurs fois."""
    previous = paragraph._p.getprevious()
    if is_h1_top_gap_spacer_element(previous):
        parent = previous.getparent()
        if parent is not None:
            parent.remove(previous)


def paragraph_has_section_break(p_elm) -> bool:
    """Retourne True si le paragraphe porte déjà un saut de section."""
    if p_elm is None or p_elm.tag != qn("w:p"):
        return False
    p_pr = p_elm.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:sectPr")) is not None


def _section_template_from_document(paragraph):
    """Copie les propriétés de section existantes pour créer un break homogène.

    Dans Word, un saut de section Next Page est porté par le paragraphe qui
    précède la nouvelle section. On copie donc le sectPr courant pour garder
    la même taille de page, les mêmes marges, headers/footers et réglages KDP.
    """
    body = paragraph._p.getparent()
    sect_pr = body.find(qn("w:sectPr")) if body is not None else None
    if sect_pr is None:
        # Fallback très rare : on crée au moins un sectPr valide.
        sect_pr = OxmlElement("w:sectPr")
    return deepcopy(sect_pr)


def insert_next_page_section_break_before(paragraph) -> bool:
    """Insère un vrai Section Break (Next Page) avant le paragraphe donné.

    Cette méthode évite les paragraphes vides et contourne le comportement de
    Word qui ignore souvent `space_before` sur un H1 placé après un simple
    page break. Le break est ajouté au paragraphe précédent, comme Word le fait
    nativement. Retourne True si un break a été ajouté.
    """
    previous = paragraph._p.getprevious()
    if previous is None or previous.tag != qn("w:p"):
        return False

    # Ne pas empiler des sauts de section si le script est relancé.
    if paragraph_has_section_break(previous):
        return False

    p_pr = previous.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        previous.insert(0, p_pr)

    sect_pr = _section_template_from_document(paragraph)
    sect_type = sect_pr.find(qn("w:type"))
    if sect_type is None:
        sect_type = OxmlElement("w:type")
        sect_pr.insert(0, sect_type)
    sect_type.set(qn("w:val"), "nextPage")

    p_pr.append(sect_pr)
    return True


def remove_direct_page_break_before(paragraph) -> None:
    """Supprime le pageBreakBefore direct pour laisser le section break gérer la page."""
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return
    for element in list(p_pr.findall(qn("w:pageBreakBefore"))):
        p_pr.remove(element)


def insert_h1_page_spacer_before(paragraph, gap_pt: float = 24.0, add_page_break: bool = True) -> None:
    """Prépare un H1 de chapitre sans faux retour de ligne.

    Pour que l'espace au-dessus du H1 soit respecté par Word, on utilise un
    vrai Section Break (Next Page) avant le H1, puis un `space_before` normal
    sur le H1. Aucun paragraphe vide n'est inséré.
    """
    remove_existing_h1_page_spacer_before(paragraph)
    remove_direct_page_break_before(paragraph)
    paragraph.paragraph_format.space_before = Pt(gap_pt)
    if add_page_break:
        insert_next_page_section_break_before(paragraph)


def is_toc_heading_text(text: str) -> bool:
    return clean_markdown_artifacts(text).strip().lower() in {
        "table des matières",
        "table of contents",
        "sommaire",
    }


def enforce_h1_section_breaks(doc: DocxDocument, cfg: PremiumConfig) -> None:
    """Garantit les Section Breaks avant les H1, même après insertion de la TOC."""
    if not cfg.chapter_page_breaks:
        return

    seen_content_before_first_h1 = False
    first_real_h1_seen = False

    for p in doc.paragraphs:
        text = p.text.strip()
        level = heading_level(p)

        if level == 1 and not is_toc_heading_text(text):
            apply_heading_paragraph_format(p, cfg, 1)
            needs_break = first_real_h1_seen or seen_content_before_first_h1
            insert_h1_page_spacer_before(
                p,
                gap_pt=cfg.h1_new_page_top_gap_pt,
                add_page_break=needs_break,
            )
            first_real_h1_seen = True
            continue

        if text and not is_toc_heading_text(text):
            seen_content_before_first_h1 = True


def standardize_paragraphs(doc: DocxDocument, cfg: PremiumConfig, stats: ProcessStats) -> None:
    first_h1_seen = False
    seen_content_before_first_h1 = False
    normal = style_lookup(doc, "Normal", "normal")
    formula_style = style_lookup(doc, "KK - Formule")
    transformation_style = style_lookup(doc, "KK - Transformation")
    code_style = style_lookup(doc, "KK - Code")
    callout_style = style_lookup(doc, "KK - Encadre")

    # Liste figée : on peut insérer des paragraphes-espace sans perturber la boucle.
    paragraphs = list(doc.paragraphs)
    for idx, p in enumerate(paragraphs):
        text = p.text.strip()

        if cfg.clear_google_run_formatting:
            for run in p.runs:
                clear_run_formatting_keep_emphasis(run)

        if cfg.normalize_heading_levels:
            level = classify_heading(text)
            if level:
                strip_markdown_heading_marks(p)
                target = style_lookup(doc, level, f"Titre {level[-1]}")
                if target is not None:
                    p.style = target
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                stats.headings += 1

                if level != "Heading 1":
                    if text and not first_h1_seen:
                        seen_content_before_first_h1 = True
                    continue

        # Titres existants Word/Titre : préserver les vrais espacements de titres
        # au lieu de les traiter comme du corps. Aucun paragraphe vide n'est
        # inséré pour créer de l'espace.
        current_heading_level = heading_level(p)
        if current_heading_level is not None:
            apply_heading_paragraph_format(p, cfg, current_heading_level)
            if current_heading_level == 1:
                if cfg.chapter_page_breaks:
                    needs_page_break = first_h1_seen or seen_content_before_first_h1
                    insert_h1_page_spacer_before(
                        p,
                        gap_pt=cfg.h1_new_page_top_gap_pt,
                        add_page_break=needs_page_break,
                    )
                else:
                    remove_direct_page_break_before(p)
                    p.paragraph_format.space_before = Pt(cfg.h1_new_page_top_gap_pt)
                first_h1_seen = True
            elif current_heading_level == 2 and cfg.page_breaks_before_h2:
                insert_manual_page_break_before(p)
            continue

        if text and not first_h1_seen:
            seen_content_before_first_h1 = True

        # Listes : compactes, jamais justifiées, et en police corps.
        if is_list_paragraph(p):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            force_list_paragraph_font(p, cfg, get_numbering_level_format(p))

            # L'espace visible avant une liste vient souvent du paragraphe
            # qui précède, pas de la liste elle-même. On le supprime au début
            # d'un bloc de liste, sans toucher aux titres.
            prev = previous_non_empty_paragraph(paragraphs, idx)
            if prev is not None and not is_list_paragraph(prev) and not previous_paragraph_is_heading(prev):
                prev.paragraph_format.space_after = Pt(0)
            continue

        if text and CALLOUT_RE.match(clean_markdown_artifacts(text)):
            if callout_style is not None:
                p.style = callout_style
            if cfg.add_callout_boxes:
                set_paragraph_shading(p, UCKK_PARCHMENT)
                set_paragraph_border(p, "left", UCKK_GOLD, "14")
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats.callouts += 1
            continue

        if text and looks_like_formula_block(text):
            if "→" in text and transformation_style is not None:
                p.style = transformation_style
                # Ne pas centrer les trajectoires/listes avec flèches.
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif formula_style is not None:
                p.style = formula_style
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats.formula_blocks += 1
            continue

        if text and looks_like_code_line(text):
            if code_style is not None:
                p.style = code_style
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats.code_lines += 1
            continue

        if normal is not None and p.style is not None:
            if p.style.name.lower() in {"normal", "normal text", "body text"}:
                p.style = normal

        # Corps courant : espace contrôlé par Word entre paragraphes normaux.
        # Aucun paragraphe vide / retour manuel n'est inséré pour créer cet espace.
        if text:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(cfg.body_space_after_pt)
            force_paragraph_runs_font(p, cfg)

            # Anti-spread : les paragraphes courts, les lignes manuelles et les
            # blocs semi-formulaires ne doivent jamais être justifiés.
            if cfg.body_alignment == "left":
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif cfg.body_alignment == "smart-justify":
                if has_manual_line_breaks(p) or is_spread_risk_text(text):
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif cfg.body_alignment == "justify":
                if has_manual_line_breaks(p) and is_spread_risk_text(text):
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def next_non_empty_paragraph(paragraphs: list, index: int):
    for j in range(index + 1, len(paragraphs)):
        candidate = paragraphs[j]
        if candidate.text.strip():
            return candidate
    return None


def add_space_after_list_blocks(doc: DocxDocument, cfg: PremiumConfig) -> None:
    """Ajoute un petit espace après le dernier item d'un bloc de liste.

    Les items internes restent compacts. Seul le dernier item avant un paragraphe
    normal reçoit un `space_after`, sans insérer de retour vide.
    """
    paragraphs = list(doc.paragraphs)
    for idx, p in enumerate(paragraphs):
        if not is_list_paragraph(p):
            continue
        nxt = next_non_empty_paragraph(paragraphs, idx)
        if nxt is None or is_list_paragraph(nxt):
            continue
        if heading_level(nxt) is not None:
            continue
        p.paragraph_format.space_after = Pt(cfg.list_block_space_after_pt)


def remove_excess_blank_paragraphs(doc: DocxDocument, stats: ProcessStats, max_blanks: int = 1) -> None:
    blanks = 0
    for p in list(doc.paragraphs):
        if p.text.strip():
            blanks = 0
            continue

        # On garde seulement les paragraphes qui portent un vrai saut de page
        # explicite. Les anciens paragraphes-espace H1 basés sur pageBreakBefore
        # sont supprimés : l'espace doit venir des propriétés Word du titre.
        if paragraph_has_explicit_page_break_run(p._p):
            blanks = 0
            continue

        blanks += 1
        if blanks > max_blanks:
            remove_paragraph(p)
            stats.blank_paragraphs_removed += 1


# ---------------------------------------------------------------------------
# Page, header, footer, tables, TOC
# ---------------------------------------------------------------------------

def disable_even_odd_headers(doc: DocxDocument) -> None:
    """Désactive les en-têtes/pieds différents pages paires/impaires.

    Le modèle de référence n'utilise pas w:evenAndOddHeaders. On retire donc
    ce réglage lorsqu'un document source l'a déjà activé.
    """
    settings = doc.settings._element
    even_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_odd is not None:
        settings.remove(even_odd)


def _set_section_common_flags(section) -> None:
    section.start_type = WD_SECTION_START.NEW_PAGE
    try:
        section.gutter = Twips(KDP_6X9_GUTTER_TWIPS)
    except Exception:
        pass
    try:
        section.different_first_page_header_footer = False
    except Exception:
        pass


def configure_sections(doc: DocxDocument, cfg: PremiumConfig) -> None:
    """Standardise toutes les sections.

    Défaut : le layout dominant du modèle KDP 6 x 9 po fourni, observé dans
    la majorité des sections de chapitres. On force toutes les sections au
    même réglage pour éviter les écarts de fin de document, préambule, etc.
    """
    if cfg.margin_mode == "recto-verso":
        enable_mirror_margins(doc)
        disable_even_odd_headers(doc)
    else:
        disable_even_odd_headers(doc)

    for section in doc.sections:
        _set_section_common_flags(section)

        if cfg.page_size.lower() in {"kdp_6x9", "6x9", "kdp-6x9"}:
            # Valeurs exactes du layout dominant du document modèle.
            section.page_width = Twips(KDP_6X9_WIDTH_TWIPS)
            section.page_height = Twips(KDP_6X9_HEIGHT_TWIPS)
            section.top_margin = Twips(KDP_6X9_TOP_TWIPS)
            section.bottom_margin = Twips(KDP_6X9_BOTTOM_TWIPS)
            section.left_margin = Twips(KDP_6X9_INSIDE_TWIPS)
            section.right_margin = Twips(KDP_6X9_OUTSIDE_TWIPS)
            section.header_distance = Twips(KDP_6X9_HEADER_TWIPS)
            section.footer_distance = Twips(KDP_6X9_FOOTER_TWIPS)
            continue

        if cfg.margin_mode == "recto-verso":
            left = cfg.margin_inside_cm
            right = cfg.margin_outside_cm
        elif cfg.margin_mode == "left":
            left = cfg.margin_inside_cm
            right = cfg.margin_outside_cm
        else:
            left = cfg.margin_left_cm
            right = cfg.margin_right_cm

        if cfg.page_size.lower() == "a4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        else:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        section.top_margin = Cm(cfg.margin_top_cm)
        section.bottom_margin = Cm(cfg.margin_bottom_cm)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)
        section.header_distance = Cm(cfg.header_distance_cm)
        section.footer_distance = Cm(cfg.footer_distance_cm)

def get_document_title(doc: DocxDocument, fallback: str = "Manuel King Klown") -> str:
    core_title = (doc.core_properties.title or "").strip()
    if core_title and core_title.lower() not in GENERIC_TITLES:
        return core_title[:90]

    for p in doc.paragraphs:
        t = clean_markdown_artifacts(p.text)
        if not t:
            continue
        if t.lower() in GENERIC_TITLES:
            continue
        if len(t) > 90:
            return t[:90].rstrip() + "…"
        return t

    return fallback


def get_heading_1_style_name(doc: DocxDocument) -> str:
    """Retourne le nom réel du style Heading 1 / Titre 1.

    Le champ Word STYLEREF a besoin du nom du style pour afficher
    automatiquement le titre du chapitre courant dans le header.
    """
    style = style_lookup(doc, "Heading 1", "Titre 1")
    if style is not None:
        return style.name
    return "Heading 1"


def add_current_chapter_field(paragraph, doc: DocxDocument, placeholder: str = "Chapitre courant") -> None:
    """Insère un champ Word dynamique qui affiche le dernier Heading 1 actif.

    Dans Word/LibreOffice, le champ STYLEREF se met à jour à l'ouverture
    ou lors de l'export PDF. Il permet d'avoir un header qui suit les chapitres
    sans créer manuellement une section par chapitre.
    """
    style_name = get_heading_1_style_name(doc)
    safe_style_name = style_name.replace('"', r'"')
    add_field(paragraph, f'STYLEREF "{safe_style_name}" \\* MERGEFORMAT', placeholder)


def format_run(run, cfg: PremiumConfig, size: float = 8.5, color: str = BLACK, bold: bool = False) -> None:
    run.font.name = cfg.heading_font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(_hex(color))
    run.font.bold = bold


def section_index_for_paragraphs(doc: DocxDocument) -> dict[int, int]:
    """Mappe chaque paragraphe vers son index de section Word.

    Dans OOXML, un Section Break est porté par le paragraphe qui termine la
    section. Le paragraphe suivant appartient donc à la section suivante.
    """
    mapping: dict[int, int] = {}
    section_index = 0
    for p in doc.paragraphs:
        mapping[id(p._p)] = section_index
        if paragraph_has_section_break(p._p):
            section_index += 1
    return mapping


def first_body_section_index(doc: DocxDocument) -> int:
    """Retourne la section où commence le premier vrai H1 du corps.

    Cela permet de ne pas numéroter la TOC/préambule et de recommencer à 1
    au premier chapitre. Si aucun H1 n'est trouvé, on part de la première section.
    """
    mapping = section_index_for_paragraphs(doc)
    for p in doc.paragraphs:
        text = p.text.strip()
        if heading_level(p) == 1 and not is_toc_heading_text(text):
            return mapping.get(id(p._p), 0)
    return 0


def clear_header_footer_part(part) -> None:
    """Vide un header/footer sans laisser de retrait ou de contenu hérité."""
    paragraph = part.paragraphs[0] if part.paragraphs else part.add_paragraph()
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def set_section_page_start(section, start: Optional[int]) -> None:
    """Définit le départ de numérotation d'une section.

    start=1 force le redémarrage à 1. start=None retire w:start pour continuer
    depuis la section précédente.
    """
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        if start is None:
            return
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)

    if start is None:
        start_attr = qn("w:start")
        if start_attr in pg_num.attrib:
            del pg_num.attrib[start_attr]
    else:
        pg_num.set(qn("w:start"), str(start))


def add_header_footer(doc: DocxDocument, cfg: PremiumConfig) -> None:
    title = get_document_title(doc)
    body_start_idx = first_body_section_index(doc)

    for section_idx, section in enumerate(doc.sections):
        # Chaque section doit être indépendante, sinon Word peut réutiliser le
        # footer de la TOC ou réinitialiser visuellement les champs.
        try:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
        except Exception:
            pass

        # Numérotation : pas de footer avant le corps; première section du corps
        # à Page 1; toutes les sections suivantes continuent automatiquement.
        if section_idx < body_start_idx:
            set_section_page_start(section, None)
            clear_header_footer_part(section.header)
            clear_header_footer_part(section.footer)
            continue
        elif section_idx == body_start_idx:
            set_section_page_start(section, 1)
        else:
            set_section_page_start(section, None)

        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = ""

        if cfg.header_layout == "none":
            pass

        elif cfg.header_layout == "chapter":
            # Recommandé : affiche automatiquement le Heading 1 courant.
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_current_chapter_field(p, doc, title)
            for run in p.runs:
                format_run(run, cfg, size=8.2, color=BLACK, bold=False)

        elif cfg.header_layout == "brand-chapter":
            # Variante : marque + chapitre courant.
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            brand = p.add_run("Univers-Cité King Klown")
            format_run(brand, cfg, size=8.2, color=BLACK, bold=True)
            sep = p.add_run(" — ")
            format_run(sep, cfg, size=8.2, color=BLACK, bold=False)
            add_current_chapter_field(p, doc, title)
            for run in p.runs:
                if run.font.size is None:
                    format_run(run, cfg, size=8.2, color=BLACK, bold=False)

        elif cfg.header_layout == "split":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            left = p.add_run("Univers-Cité King Klown")
            format_run(left, cfg, size=8.2, color=BLACK, bold=True)
            p.add_run("\t")
            add_current_chapter_field(p, doc, title)
            for run in p.runs:
                if run.font.size is None:
                    format_run(run, cfg, size=8.2, color=BLACK, bold=False)

        elif cfg.header_layout == "minimal":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("King Klown")
            format_run(r, cfg, size=8.2, color=BLACK, bold=True)

        else:
            # Header statique, surtout utile pour les documents courts.
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(title)
            format_run(r, cfg, size=8.4, color=BLACK, bold=False)

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fp.paragraph_format.left_indent = Pt(0)
        fp.paragraph_format.first_line_indent = Pt(0)
        fp.paragraph_format.right_indent = Pt(0)
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)

        # Format demandé : pas de tiret, pas d'indentation, 7 tabulations.
        left = fp.add_run(cfg.footer_text + ("\t" * 7) + "Page ")
        format_run(left, cfg, size=8.2, color=BLACK)

        # Affiche seulement le numéro de page courant, sans total du type " / 252".
        add_field(fp, "PAGE", "1")

        for run in fp.runs:
            if run.font.size is None:
                format_run(run, cfg, size=8.2, color=BLACK)

def normalize_toc_field_instruction(instr: str) -> str:
    """Retourne une instruction TOC limitée strictement aux niveaux 1-2.

    On retire le commutateur TOC backslash-u, parce qu'il peut reprendre des
    niveaux d'outline du document et laisser revenir des H3 dans la table.
    """
    if "TOC" not in instr.upper():
        return instr

    updated = re.sub(r'\\o\s+"[^"]+"', r'\\o "1-2"', instr)
    if updated == instr and "\\o" not in instr:
        updated = re.sub(r'\bTOC\b', r'TOC \\o "1-2"', instr, count=1, flags=re.IGNORECASE)

    # Enlève le commutateur TOC backslash-u, avec ou sans espaces autour.
    updated = re.sub(r'\s+\\u\b', '', updated)
    updated = re.sub(r'\\u\b', '', updated)
    return re.sub(r'\s{2,}', ' ', updated).strip()


def force_toc_two_levels(doc: DocxDocument) -> bool:
    """Force les champs TOC existants à n'afficher que Heading 1 et 2.

    Le script insérait déjà une TOC en 1-2, mais si le document contenait
    une ancienne TOC en 1-3, l'ancien champ restait intact. Cette fonction
    modifie les champs existants, y compris ceux créés par Word.
    """
    changed = False

    for paragraph in doc.paragraphs:
        for instr in paragraph._p.findall('.//' + qn('w:instrText')):
            if instr.text and "TOC" in instr.text.upper():
                new_text = normalize_toc_field_instruction(instr.text)
                if new_text != instr.text:
                    instr.text = new_text
                    changed = True

        for simple in paragraph._p.findall('.//' + qn('w:fldSimple')):
            instr_value = simple.get(qn('w:instr'))
            if instr_value and "TOC" in instr_value.upper():
                new_value = normalize_toc_field_instruction(instr_value)
                if new_value != instr_value:
                    simple.set(qn('w:instr'), new_value)
                    changed = True

    return changed


def remove_visible_toc_entries_deeper_than_two(doc: DocxDocument) -> int:
    """Supprime les lignes visibles TOC 3+ laissées par une ancienne TOC.

    Word régénèrera la table à l'ouverture grâce à updateFields, mais cette
    suppression évite de voir encore des entrées niveau 3 dans le résultat
    avant l'actualisation manuelle.
    """
    removed = 0
    for paragraph in list(doc.paragraphs):
        style_name = getattr(paragraph.style, "name", "") or ""
        style_lower = style_name.strip().lower()
        if re.search(r'\b(toc|tm)\s*[3-9]\b', style_lower) or re.search(r'table des matières\s*[3-9]', style_lower):
            remove_paragraph(paragraph)
            removed += 1
    return removed


def insert_toc_at_start(doc: DocxDocument) -> None:
    if not doc.paragraphs:
        doc.add_paragraph()

    force_toc_two_levels(doc)
    remove_visible_toc_entries_deeper_than_two(doc)

    first = doc.paragraphs[0]

    # Si une TOC existe déjà au début, on la garde mais son champ est forcé à 1-2.
    for p in doc.paragraphs[:20]:
        if p.text.strip().lower() in {"table des matières", "table of contents", "sommaire"}:
            return

    toc_title = first.insert_paragraph_before("Table des matières", style="Heading 1")
    toc_title.paragraph_format.page_break_before = False

    toc_p = first.insert_paragraph_before("")
    add_field(toc_p, r'TOC \o "1-2" \h \z', "Cliquez ici puis actualisez la table des matières.")

    sep = first.insert_paragraph_before("")
    sep.add_run().add_break(WD_BREAK.PAGE)

def autofit_table_to_window(table) -> None:
    """Applique l'équivalent Word de AutoFit > AutoFit Window.

    python-docx expose `table.autofit = True`, mais ce réglage seul ne force
    pas toujours la largeur du tableau à 100 % de la zone texte dans Word.
    On ajoute donc aussi une largeur préférée en pourcentage OOXML : 5000 = 100 %.
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    tbl_pr = table._tbl.tblPr

    tbl_w = get_or_add(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    tbl_layout = get_or_add(tbl_pr, "w:tblLayout")
    tbl_layout.set(qn("w:type"), "autofit")

    # Une indentation de tableau héritée peut empêcher le vrai plein-largeur.
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is not None:
        tbl_pr.remove(tbl_ind)


def standardize_tables(doc: DocxDocument, cfg: PremiumConfig, stats: ProcessStats) -> None:
    normal = style_lookup(doc, "Normal", "normal")
    stats.tables = len(doc.tables)

    for table in doc.tables:
        autofit_table_to_window(table)
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_borders(cell)
                set_cell_margins(cell)

                if row_idx == 0:
                    set_cell_shading(cell, BLACK)
                elif row_idx % 2 == 0:
                    set_cell_shading(cell, UCKK_LIGHT_GRAY)
                else:
                    set_cell_shading(cell, WHITE)

                for p in cell.paragraphs:
                    if normal is not None:
                        p.style = normal
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for run in p.runs:
                        if cfg.clear_google_run_formatting:
                            clear_run_formatting_keep_emphasis(run)
                        run.font.name = cfg.body_font
                        run.font.size = Pt(9.7)
                        if row_idx == 0:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor.from_string(WHITE)
                        else:
                            run.font.color.rgb = RGBColor.from_string(BLACK)


def mark_header_rows_repeat(doc: DocxDocument) -> None:
    for table in doc.tables:
        if not table.rows:
            continue
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = tr_pr.find(qn("w:tblHeader"))
        if tbl_header is None:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------

def backup_file(path: Path) -> Path:
    stamp = time.strftime("%Y%m%d_%H%M%S")
    backup = path.with_suffix(path.suffix + f".bak_{stamp}")
    backup.write_bytes(path.read_bytes())
    return backup


def validate_input_path(input_path: Path) -> None:
    if not input_path.exists():
        raise FileNotFoundError(f"Introuvable : {input_path}")
    if input_path.is_file() and input_path.suffix.lower() != ".docx":
        raise ValueError(f"Le fichier n'est pas un .docx : {input_path}")


def process_docx(input_path: Path, output_path: Path, cfg: PremiumConfig,
                 dry_run: bool = False, in_place: bool = False) -> ProcessStats:
    validate_input_path(input_path)

    stats = ProcessStats(source=input_path, output=output_path, dry_run=dry_run)
    doc = Document(str(input_path))

    configure_sections(doc, cfg)
    configure_styles(doc, cfg)
    standardize_numbering_fonts(doc, cfg)
    standardize_paragraphs(doc, cfg, stats)
    # Aucun retour de ligne vide n'est conservé pour créer les espacements ;
    # l'espacement est géré par les styles/paragraph_format Word.
    remove_excess_blank_paragraphs(doc, stats, max_blanks=0)
    add_space_after_list_blocks(doc, cfg)

    if cfg.standardize_tables:
        standardize_tables(doc, cfg, stats)
        mark_header_rows_repeat(doc)

    # Si une TOC existe déjà, on la force aussi en 2 niveaux même quand
    # l'option d'ajout de TOC n'est pas cochée.
    force_toc_two_levels(doc)
    remove_visible_toc_entries_deeper_than_two(doc)

    if cfg.add_toc:
        insert_toc_at_start(doc)

    # Les Section Breaks doivent être finalisés avant les headers/footers,
    # sinon les nouvelles sections héritent mal des pieds de page et de la
    # numérotation.
    enforce_h1_section_breaks(doc, cfg)
    configure_sections(doc, cfg)

    if cfg.add_header_footer:
        add_header_footer(doc, cfg)

    if dry_run:
        return stats

    output_path.parent.mkdir(parents=True, exist_ok=True)

    if in_place:
        backup = backup_file(input_path)
        stats.warnings.append(f"Sauvegarde créée : {backup.name}")

    doc.save(str(output_path))

    if cfg.add_toc or cfg.header_layout in {"chapter", "brand-chapter", "split"}:
        try:
            enable_update_fields_on_open(output_path)
        except Exception as exc:
            stats.warnings.append(f"Champs Word : impossible d'activer updateFields : {exc}")

    if cfg.export_pdf:
        try:
            pdf_path = export_to_pdf(output_path)
            stats.warnings.append(f"PDF exporté : {pdf_path.name}")
        except Exception as exc:
            stats.warnings.append(f"Export PDF ignoré : {exc}")

    return stats


def iter_docx_files(path: Path, recursive: bool) -> Iterable[Path]:
    if path.is_file():
        if path.suffix.lower() == ".docx" and not path.name.startswith("~$"):
            yield path
        return

    pattern = "**/*.docx" if recursive else "*.docx"
    for p in sorted(path.glob(pattern)):
        if p.is_file() and not p.name.startswith("~$"):
            yield p


def output_for(input_file: Path, input_root: Path, output_root: Path, suffix: str, in_place: bool) -> Path:
    if in_place:
        return input_file

    if input_root.is_file():
        rel = Path(input_file.name)
    else:
        rel = input_file.relative_to(input_root)

    out = output_root / rel
    return out.with_name(out.stem + suffix + out.suffix)


def find_libreoffice() -> Optional[str]:
    for candidate in ("soffice", "libreoffice"):
        found = shutil.which(candidate)
        if found:
            return found

    win_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for p in win_paths:
        if Path(p).exists():
            return p

    return None


def export_to_pdf(docx_path: Path) -> Path:
    soffice = find_libreoffice()
    if not soffice:
        raise RuntimeError("LibreOffice introuvable.")

    out_dir = docx_path.parent
    cmd = [
        soffice, "--headless", "--convert-to", "pdf",
        "--outdir", str(out_dir), str(docx_path),
    ]
    completed = subprocess.run(cmd, capture_output=True, text=True, check=False)
    if completed.returncode != 0:
        raise RuntimeError(completed.stderr.strip() or completed.stdout.strip() or "Échec LibreOffice")

    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError("PDF attendu non créé.")
    return pdf_path


def run_batch(input_root: Path, output_root: Path, cfg: PremiumConfig,
              recursive: bool = False, in_place: bool = False,
              suffix: str = "UCKK_formated", dry_run: bool = False,
              logger: Optional[Callable[[str], None]] = None) -> tuple[list[ProcessStats], int]:
    log = logger or print
    validate_input_path(input_root)

    files = list(iter_docx_files(input_root, recursive))
    if not files:
        log("Aucun fichier .docx trouvé.")
        return [], 0

    log(f"{len(files)} fichier(s) à traiter.")
    stats_list: list[ProcessStats] = []
    errors = 0

    for input_file in files:
        out = output_for(input_file, input_root, output_root, suffix, in_place)
        try:
            stats = process_docx(input_file, out, cfg, dry_run=dry_run, in_place=in_place)
            stats_list.append(stats)
            prefix = "DRY-RUN" if dry_run else "OK"
            log(
                f"{prefix}  {input_file.name} → {out}\n"
                f"      titres={stats.headings}, encadrés={stats.callouts}, "
                f"formules={stats.formula_blocks}, code={stats.code_lines}, "
                f"tableaux={stats.tables}, blancs_supprimés={stats.blank_paragraphs_removed}"
            )
            for warning in stats.warnings:
                log(f"      note: {warning}")
        except Exception as exc:
            errors += 1
            log(f"ERREUR  {input_file} : {exc}")

    log(f"Terminé avec {errors} erreur(s)." if errors else "Terminé.")
    return stats_list, errors


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args(argv: Optional[list[str]] = None):
    parser = argparse.ArgumentParser(description="Standardisation premium King Klown / UCKK pour .docx.")
    parser.add_argument("input", nargs="?", help="Fichier .docx ou dossier. Si absent : ouvre le GUI.")
    parser.add_argument("--output", "-o", default=str(DEFAULT_OUTPUT_DIR), help="Dossier de sortie si --in-place n'est pas utilisé")
    parser.add_argument("--recursive", "-r", action="store_true", help="Traiter aussi les sous-dossiers")
    parser.add_argument("--in-place", action="store_true", help="Écraser les originaux avec backup automatique")
    parser.add_argument("--suffix", default=DEFAULT_SUFFIX, help="Suffixe ajouté aux fichiers générés")

    parser.add_argument("--page-size", choices=["kdp_6x9", "letter", "a4"], default=DEFAULT_PAGE_SIZE)
    parser.add_argument("--margin-mode", choices=["recto-verso", "centered", "left"], default=DEFAULT_MARGIN_MODE)
    parser.add_argument("--header-layout", choices=["chapter", "brand-chapter", "centered", "split", "minimal", "none"], default=DEFAULT_HEADER_LAYOUT, help="Header. Recommandé : chapter pour afficher le chapitre courant.")
    parser.add_argument("--body-alignment", choices=["left", "smart-justify", "justify"], default=DEFAULT_BODY_ALIGNMENT, help="Alignement du corps. Recommandé : left pour éviter les blancs étirés.")

    parser.add_argument("--toc", action="store_true", help="Insérer une table des matières")
    parser.add_argument("--no-header-footer", action="store_true", help="Ne pas ajouter header/footer")
    parser.add_argument("--no-chapter-breaks", action="store_true", help="Ne pas forcer les chapitres en nouvelle page")
    parser.add_argument("--h1-top-gap-pt", type=float, default=DEFAULT_H1_TOP_GAP_PT, help="Espace réel avant un H1 après saut de section Next Page")
    parser.add_argument("--h2-page-breaks", action="store_true", default=DEFAULT_H2_PAGE_BREAKS, help="Insérer un saut de page manuel avant chaque Heading 2 / Titre 2")
    parser.add_argument("--list-after-space-pt", type=float, default=DEFAULT_LIST_BLOCK_SPACE_AFTER_PT, help="Petit espace après le dernier item d’un bloc de liste")
    parser.add_argument("--body-space-after-pt", type=float, default=DEFAULT_BODY_SPACE_AFTER_PT, help="Espace après les paragraphes de texte normal, sans retour de ligne vide")
    parser.add_argument("--heading-levels", action="store_true", help="Détecter/changer automatiquement les niveaux Heading 1-4")
    parser.add_argument("--no-heading-levels", action="store_true", help="Compatibilité : ne pas détecter/changer automatiquement les niveaux Heading 1-4")
    parser.add_argument("--keep-google-formatting", action="store_true", help="Conserver le formatage direct Google Docs")
    parser.add_argument("--no-tables", action="store_true", help="Ne pas standardiser les tableaux")
    parser.add_argument("--callout-boxes", action="store_true", help="Encadrer visuellement les blocs pédagogiques")
    parser.add_argument("--no-callout-boxes", action="store_true", help="Compatibilité : ne pas encadrer les blocs pédagogiques")
    parser.add_argument("--body-font", default=DEFAULT_BODY_FONT)
    parser.add_argument("--heading-font", default=DEFAULT_HEADING_FONT)
    parser.add_argument("--footer-text", default=DEFAULT_FOOTER_TEXT)
    parser.add_argument("--dry-run", action="store_true", help="Analyser sans sauvegarder")
    parser.add_argument("--pdf", action="store_true", help="Exporter aussi PDF via LibreOffice")
    return parser.parse_args(argv)


def config_from_args(args) -> PremiumConfig:
    cfg = PremiumConfig()
    cfg.page_size = args.page_size
    cfg.margin_mode = args.margin_mode
    cfg.header_layout = args.header_layout
    cfg.body_alignment = args.body_alignment
    cfg.add_toc = args.toc
    cfg.add_header_footer = not args.no_header_footer
    cfg.chapter_page_breaks = not args.no_chapter_breaks
    cfg.h1_new_page_top_gap_pt = args.h1_top_gap_pt
    cfg.page_breaks_before_h2 = args.h2_page_breaks
    cfg.list_block_space_after_pt = args.list_after_space_pt
    cfg.body_space_after_pt = args.body_space_after_pt
    cfg.normalize_heading_levels = bool(args.heading_levels) and not args.no_heading_levels
    cfg.clear_google_run_formatting = not args.keep_google_formatting
    cfg.standardize_tables = not args.no_tables
    cfg.add_callout_boxes = bool(args.callout_boxes) and not args.no_callout_boxes
    cfg.body_font = args.body_font
    cfg.heading_font = args.heading_font
    cfg.footer_text = args.footer_text
    cfg.export_pdf = args.pdf
    return cfg


def cli_main(args) -> int:
    if not args.input:
        launch_gui()
        return 0

    input_root = Path(args.input).expanduser().resolve()
    output_root = Path(args.output).expanduser().resolve()
    cfg = config_from_args(args)

    _, errors = run_batch(
        input_root=input_root,
        output_root=output_root,
        cfg=cfg,
        recursive=args.recursive,
        in_place=args.in_place,
        suffix=args.suffix,
        dry_run=args.dry_run,
    )
    return 1 if errors else 0


# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------

class PremiumGUI(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("King Klown Premium Builder")
        self.geometry("880x640")
        self.minsize(780, 580)

        self.log_queue: queue.Queue[object] = queue.Queue()
        self.worker_thread: Optional[threading.Thread] = None
        self.last_converted_file: Optional[Path] = None

        self.input_path = tk.StringVar()
        self.output_path = tk.StringVar(value=str(DEFAULT_OUTPUT_DIR))

        self.suffix = tk.StringVar(value=DEFAULT_SUFFIX)
        self.body_font = tk.StringVar(value=DEFAULT_BODY_FONT)
        self.heading_font = tk.StringVar(value=DEFAULT_HEADING_FONT)
        self.footer_text = tk.StringVar(value=DEFAULT_FOOTER_TEXT)

        self.page_size = tk.StringVar(value=DEFAULT_PAGE_SIZE)
        self.margin_mode = tk.StringVar(value=DEFAULT_MARGIN_MODE)
        self.header_layout = tk.StringVar(value=DEFAULT_HEADER_LAYOUT)
        self.body_alignment = tk.StringVar(value=DEFAULT_BODY_ALIGNMENT)

        self.recursive = tk.BooleanVar(value=DEFAULT_RECURSIVE)
        self.in_place = tk.BooleanVar(value=DEFAULT_IN_PLACE)
        self.add_toc = tk.BooleanVar(value=DEFAULT_ADD_TOC_GUI)
        self.header_footer = tk.BooleanVar(value=DEFAULT_ADD_HEADER_FOOTER)
        self.chapter_breaks = tk.BooleanVar(value=DEFAULT_CHAPTER_PAGE_BREAKS)
        self.h2_page_breaks = tk.BooleanVar(value=DEFAULT_H2_PAGE_BREAKS)
        self.do_not_change_heading_levels = tk.BooleanVar(value=DEFAULT_DO_NOT_CHANGE_HEADING_LEVELS)
        self.clear_google_formatting = tk.BooleanVar(value=DEFAULT_CLEAR_GOOGLE_FORMATTING)
        self.standardize_tables_var = tk.BooleanVar(value=DEFAULT_STANDARDIZE_TABLES)
        self.callout_boxes = tk.BooleanVar(value=DEFAULT_CALLOUT_BOXES)
        self.dry_run = tk.BooleanVar(value=DEFAULT_DRY_RUN)
        self.export_pdf_var = tk.BooleanVar(value=DEFAULT_EXPORT_PDF)

        self._build_ui()
        self.after(120, self._pump_log_queue)

    def _build_ui(self) -> None:
        root = ttk.Frame(self, padding=12)
        root.pack(fill="both", expand=True)

        source = ttk.LabelFrame(root, text="Source")
        source.pack(fill="x", pady=(0, 10))
        ttk.Label(source, text="Fichier ou dossier .docx").grid(row=0, column=0, sticky="w", padx=8, pady=8)
        ttk.Entry(source, textvariable=self.input_path).grid(row=0, column=1, sticky="ew", padx=8, pady=8)
        ttk.Button(source, text="Fichier…", command=self._choose_file).grid(row=0, column=2, padx=4, pady=8)
        ttk.Button(source, text="Dossier…", command=self._choose_folder).grid(row=0, column=3, padx=8, pady=8)
        source.columnconfigure(1, weight=1)

        out = ttk.LabelFrame(root, text="Sortie")
        out.pack(fill="x", pady=(0, 10))
        ttk.Label(out, text="Dossier de sortie").grid(row=0, column=0, sticky="w", padx=8, pady=8)
        self.output_entry = ttk.Entry(out, textvariable=self.output_path)
        self.output_entry.grid(row=0, column=1, sticky="ew", padx=8, pady=8)
        ttk.Button(out, text="Choisir…", command=self._choose_output).grid(row=0, column=2, padx=8, pady=8)
        ttk.Label(out, text="Suffixe").grid(row=1, column=0, sticky="w", padx=8, pady=8)
        ttk.Entry(out, textvariable=self.suffix, width=18).grid(row=1, column=1, sticky="w", padx=8, pady=8)
        out.columnconfigure(1, weight=1)

        layout = ttk.LabelFrame(root, text="Layout premium")
        layout.pack(fill="x", pady=(0, 10))

        ttk.Label(layout, text="Page").grid(row=0, column=0, sticky="w", padx=8, pady=6)
        ttk.Combobox(layout, textvariable=self.page_size, values=["kdp_6x9", "letter", "a4"], state="readonly", width=14).grid(row=0, column=1, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Marges").grid(row=0, column=2, sticky="w", padx=8, pady=6)
        ttk.Combobox(
            layout,
            textvariable=self.margin_mode,
            values=["recto-verso", "centered", "left"],
            state="readonly",
            width=16
        ).grid(row=0, column=3, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Header").grid(row=0, column=4, sticky="w", padx=8, pady=6)
        ttk.Combobox(
            layout,
            textvariable=self.header_layout,
            values=["chapter", "brand-chapter", "centered", "split", "minimal", "none"],
            state="readonly",
            width=16
        ).grid(row=0, column=5, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Alignement").grid(row=1, column=4, sticky="w", padx=8, pady=6)
        ttk.Combobox(
            layout,
            textvariable=self.body_alignment,
            values=["left", "smart-justify", "justify"],
            state="readonly",
            width=14
        ).grid(row=1, column=5, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Police corps").grid(row=1, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(layout, textvariable=self.body_font, width=18).grid(row=1, column=1, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Police titres").grid(row=1, column=2, sticky="w", padx=8, pady=6)
        ttk.Entry(layout, textvariable=self.heading_font, width=18).grid(row=1, column=3, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Pied de page").grid(row=2, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(layout, textvariable=self.footer_text).grid(row=2, column=1, columnspan=5, sticky="ew", padx=8, pady=6)
        layout.columnconfigure(5, weight=1)

        checks = ttk.LabelFrame(root, text="Traitement")
        checks.pack(fill="x", pady=(0, 10))

        checkbox_data = [
            ("Récursif", self.recursive),
            ("Écraser originaux avec backup", self.in_place),
            ("Table des matières", self.add_toc),
            ("En-tête / pied de page", self.header_footer),
            ("Sauts de page chapitres", self.chapter_breaks),
            ("Page break manuel avant tous les H2", self.h2_page_breaks),
            ("Ne pas changer heading level", self.do_not_change_heading_levels),
            ("Nettoyer formatage Google Docs", self.clear_google_formatting),
            ("Standardiser tableaux", self.standardize_tables_var),
            ("Encadrés pédagogiques", self.callout_boxes),
            ("Dry-run seulement", self.dry_run),
            ("Exporter PDF via LibreOffice", self.export_pdf_var),
        ]

        for idx, (label, var) in enumerate(checkbox_data):
            r = idx // 2
            c = (idx % 2) * 2
            ttk.Checkbutton(checks, text=label, variable=var, command=self._toggle_in_place).grid(
                row=r, column=c, columnspan=2, sticky="w", padx=8, pady=4
            )

        actions = ttk.Frame(root)
        actions.pack(fill="x", pady=(0, 10))
        self.run_button = ttk.Button(actions, text="Standardiser premium", command=self._start_processing)
        self.run_button.pack(side="left")
        self.open_button = ttk.Button(
            actions,
            text="Open converted file",
            command=self._open_converted_file,
            state="disabled",
        )
        self.open_button.pack(side="left", padx=(8, 0))
        ttk.Button(actions, text="Effacer le log", command=self._clear_log).pack(side="left", padx=8)
        self.progress = ttk.Progressbar(actions, mode="indeterminate")
        self.progress.pack(side="right", fill="x", expand=True, padx=(8, 0))

        log_frame = ttk.LabelFrame(root, text="Log")
        log_frame.pack(fill="both", expand=True)
        self.log_text = tk.Text(log_frame, wrap="word", height=14)
        self.log_text.pack(side="left", fill="both", expand=True)
        scroll = ttk.Scrollbar(log_frame, command=self.log_text.yview)
        scroll.pack(side="right", fill="y")
        self.log_text.configure(yscrollcommand=scroll.set)

    def _choose_file(self) -> None:
        path = filedialog.askopenfilename(
            title="Choisir un fichier .docx",
            filetypes=[("Word documents", "*.docx"), ("Tous les fichiers", "*.*")]
        )
        if path:
            self.input_path.set(path)

    def _choose_folder(self) -> None:
        path = filedialog.askdirectory(title="Choisir un dossier")
        if path:
            self.input_path.set(path)

    def _choose_output(self) -> None:
        path = filedialog.askdirectory(title="Choisir le dossier de sortie")
        if path:
            self.output_path.set(path)

    def _toggle_in_place(self) -> None:
        self.output_entry.configure(state="disabled" if self.in_place.get() else "normal")

    def _clear_log(self) -> None:
        self.log_text.delete("1.0", "end")

    def _log(self, msg: str) -> None:
        self.log_text.insert("end", msg + "\n")
        self.log_text.see("end")

    def _thread_log(self, msg: object) -> None:
        self.log_queue.put(msg)

    def _pump_log_queue(self) -> None:
        try:
            while True:
                msg = self.log_queue.get_nowait()
                if isinstance(msg, tuple) and len(msg) == 2 and msg[0] == "__OPEN_FILE__":
                    self.last_converted_file = Path(str(msg[1]))
                    self.open_button.configure(state="normal")
                    self._log(f"Fichier converti prêt : {self.last_converted_file}")
                elif msg == "__DONE__":
                    self.progress.stop()
                    self.run_button.configure(state="normal")
                else:
                    self._log(str(msg))
        except queue.Empty:
            pass
        self.after(120, self._pump_log_queue)

    def _open_converted_file(self) -> None:
        path = self.last_converted_file
        if path is None:
            messagebox.showinfo("Aucun fichier", "Aucun fichier converti disponible à ouvrir.")
            return
        if not path.exists():
            messagebox.showerror("Fichier introuvable", f"Le fichier n'existe plus :\n{path}")
            self.open_button.configure(state="disabled")
            return

        try:
            if sys.platform.startswith("win"):
                os.startfile(str(path))  # type: ignore[attr-defined]
            elif sys.platform == "darwin":
                subprocess.Popen(["open", str(path)])
            else:
                subprocess.Popen(["xdg-open", str(path)])
        except Exception as exc:
            messagebox.showerror("Ouverture impossible", f"Impossible d'ouvrir le fichier :\n{exc}")

    def _start_processing(self) -> None:
        if self.worker_thread and self.worker_thread.is_alive():
            messagebox.showinfo("Traitement en cours", "Un traitement est déjà en cours.")
            return

        source = self.input_path.get().strip()
        if not source:
            messagebox.showerror("Source manquante", "Choisis un fichier .docx ou un dossier.")
            return

        input_root = Path(source).expanduser().resolve()
        output_root = Path(self.output_path.get().strip() or str(DEFAULT_OUTPUT_DIR)).expanduser().resolve()

        if self.in_place.get():
            confirm = messagebox.askyesno(
                "Confirmer le mode in-place",
                "Le mode in-place écrase les fichiers originaux, avec une sauvegarde .bak automatique.\n\nContinuer ?"
            )
            if not confirm:
                return

        cfg = PremiumConfig(
            page_size=self.page_size.get(),
            margin_mode=self.margin_mode.get(),
            header_layout=self.header_layout.get(),
            body_alignment=self.body_alignment.get(),
            body_font=self.body_font.get().strip() or DEFAULT_BODY_FONT,
            heading_font=self.heading_font.get().strip() or DEFAULT_HEADING_FONT,
            footer_text=self.footer_text.get().strip() or DEFAULT_FOOTER_TEXT,
            add_toc=self.add_toc.get(),
            add_header_footer=self.header_footer.get(),
            chapter_page_breaks=self.chapter_breaks.get(),
            page_breaks_before_h2=self.h2_page_breaks.get(),
            normalize_heading_levels=not self.do_not_change_heading_levels.get(),
            clear_google_run_formatting=self.clear_google_formatting.get(),
            standardize_tables=self.standardize_tables_var.get(),
            add_callout_boxes=self.callout_boxes.get(),
            export_pdf=self.export_pdf_var.get(),
        )

        self.last_converted_file = None
        self.open_button.configure(state="disabled")
        self.run_button.configure(state="disabled")
        self.progress.start(10)
        self._log("—" * 72)
        self._log(
            f"Premium | page={cfg.page_size} | marges={cfg.margin_mode} | "
            f"header={cfg.header_layout} | alignement={cfg.body_alignment} | espace corps={cfg.body_space_after_pt} pt | headings={'auto' if cfg.normalize_heading_levels else 'inchangés'} | H2 page break manuel={'oui' if cfg.page_breaks_before_h2 else 'non'} | texte=noir | corps={cfg.body_font} | titres={cfg.heading_font}"
        )

        def worker() -> None:
            try:
                stats_list, _errors = run_batch(
                    input_root=input_root,
                    output_root=output_root,
                    cfg=cfg,
                    recursive=self.recursive.get(),
                    in_place=self.in_place.get(),
                    suffix=self.suffix.get().strip() or DEFAULT_SUFFIX,
                    dry_run=self.dry_run.get(),
                    logger=self._thread_log,
                )
                if not self.dry_run.get():
                    for stats in reversed(stats_list):
                        if stats.output.exists():
                            self._thread_log(("__OPEN_FILE__", str(stats.output)))
                            break
            except Exception as exc:
                self._thread_log(f"ERREUR GÉNÉRALE : {exc}")
            finally:
                self._thread_log("__DONE__")

        self.worker_thread = threading.Thread(target=worker, daemon=True)
        self.worker_thread.start()


def launch_gui() -> None:
    app = PremiumGUI()
    app.mainloop()


def main(argv: Optional[list[str]] = None) -> int:
    args = parse_args(argv)
    return cli_main(args)


if __name__ == "__main__":
    raise SystemExit(main())
