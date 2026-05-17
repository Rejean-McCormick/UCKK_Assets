#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
UCKK DOCX Table Styler — version simple

- Sélection d'un fichier DOCX source.
- Sélection d'un fichier DOCX de sortie.
- Sélection optionnelle d'un seul PNG décoratif.
- Possibilité de styliser uniquement les tableaux, sans ajouter de PNG.
- Pas de prévisualisation.

Installation :
    pip install -r requirements.txt

Lancement :
    python uckk_docx_table_gui_simple.py
"""

from __future__ import annotations

import threading
import traceback
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, List, Optional

import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    from docx.text.paragraph import Paragraph
except Exception as exc:  # pragma: no cover
    Document = None
    IMPORT_ERROR = exc
else:
    IMPORT_ERROR = None


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

@dataclass
class UCKKConfig:
    input_docx: str = ""
    output_docx: str = ""
    png_file: str = ""

    add_png: bool = False
    color_mode: str = "BW"  # BW | COULEUR

    max_separator_width_pt: float = 310.0
    table_border_width_pt: float = 0.5

    header_font_size: float = 8.0
    body_font_size: float = 8.0
    font_title: str = "Georgia"
    font_body: str = "Arial"

    first_row_as_header: bool = True
    accent_first_column: bool = True
    alternate_rows: bool = True
    apply_column_widths: bool = True
    column_widths_pt: str = "120,175,175,175,175,175,175,175"

    skip_small_tables: bool = False
    min_rows: int = 2
    min_cols: int = 2


@dataclass
class ProcessingResult:
    tables_found: int = 0
    tables_styled: int = 0
    tables_skipped: int = 0
    separators_inserted: int = 0
    warnings: List[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Traitement DOCX
# ---------------------------------------------------------------------------

def process_docx(config: UCKKConfig, log: Optional[Callable[[str], None]] = None) -> ProcessingResult:
    if IMPORT_ERROR is not None:
        raise RuntimeError(
            "Le module python-docx n'est pas installé. "
            "Installe-le avec : pip install python-docx"
        ) from IMPORT_ERROR

    input_path = Path(config.input_docx).expanduser()
    output_path = Path(config.output_docx).expanduser()

    if not input_path.exists():
        raise FileNotFoundError(f"Fichier introuvable : {input_path}")

    if input_path.suffix.lower() != ".docx":
        raise ValueError("Le fichier d'entrée doit être un .docx.")

    if output_path.suffix.lower() != ".docx":
        output_path = output_path.with_suffix(".docx")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    png_path = resolve_png_path(config)
    palette = get_palette(config.color_mode)

    log_msg(log, f"Ouverture du document : {input_path.name}")
    doc = Document(str(input_path))
    tables = list(doc.tables)

    result = ProcessingResult(tables_found=len(tables))

    log_msg(log, f"Tableaux trouvés : {len(tables)}")
    if config.add_png:
        log_msg(log, f"PNG sélectionné : {png_path.name}")
    else:
        log_msg(log, "PNG : désactivé — stylisation des tableaux seulement")

    for idx, table in enumerate(tables, start=1):
        try:
            if not should_process_table(table, config):
                result.tables_skipped += 1
                continue

            style_table(table, palette, config)
            result.tables_styled += 1

            if config.add_png and png_path is not None:
                ok = insert_png_before_table(table, png_path, config)
                if ok:
                    result.separators_inserted += 1
                else:
                    result.warnings.append(f"Tableau {idx} : le PNG n'a pas pu être inséré.")

            if idx % 10 == 0:
                log_msg(log, f"Progression : {idx}/{len(tables)} tableaux lus")

        except Exception as exc:
            result.warnings.append(f"Tableau {idx} : {exc}")

    log_msg(log, f"Enregistrement : {output_path.name}")
    doc.save(str(output_path))

    log_msg(log, "Terminé.")
    return result


def resolve_png_path(config: UCKKConfig) -> Optional[Path]:
    if not config.add_png:
        return None

    if not config.png_file.strip():
        raise ValueError("Le mode PNG est activé, mais aucun fichier PNG n'est sélectionné.")

    path = Path(config.png_file).expanduser()

    if not path.exists():
        raise FileNotFoundError(f"PNG introuvable : {path}")

    if path.suffix.lower() != ".png":
        raise ValueError("Le séparateur doit être un fichier .png.")

    return path


def should_process_table(table, config: UCKKConfig) -> bool:
    if not config.skip_small_tables:
        return True

    rows = len(table.rows)
    cols = len(table.columns)
    return rows >= int(config.min_rows) and cols >= int(config.min_cols)


def style_table(table, palette: dict, config: UCKKConfig) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    try:
        table.autofit = False
    except Exception:
        pass

    set_table_borders(table, palette["line"], config.table_border_width_pt)
    set_table_layout_fixed(table)

    widths = parse_widths(config.column_widths_pt)

    for r, row in enumerate(table.rows):
        if config.first_row_as_header and r == 0:
            set_repeat_table_header(row)

        set_row_cant_split(row)

        for c, cell in enumerate(row.cells):
            is_header = config.first_row_as_header and r == 0
            is_first_col = config.accent_first_column and c == 0 and not is_header
            is_alt = config.alternate_rows and (r % 2 == 0) and not is_header and not is_first_col

            fill = palette["surface"]
            if is_header:
                fill = palette["petrol"]
            elif is_first_col:
                fill = palette["left_col"]
            elif is_alt:
                fill = palette["alt"]

            set_cell_shading(cell, fill)
            set_cell_borders(cell, palette["line"], config.table_border_width_pt)
            set_cell_margins(cell, top=18, bottom=18, start=58, end=58)

            try:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            except Exception:
                pass

            if config.apply_column_widths and c < len(widths):
                try:
                    cell.width = Pt(widths[c])
                except Exception:
                    pass

            style_cell_text(cell, is_header, is_first_col, palette, config)


def style_cell_text(cell, is_header: bool, is_first_col: bool, palette: dict, config: UCKKConfig) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

        for run in paragraph.runs:
            run.font.name = config.font_title if is_header else config.font_body
            run._element.rPr.rFonts.set(qn("w:eastAsia"), config.font_title if is_header else config.font_body)

            run.font.size = Pt(config.header_font_size if is_header else config.body_font_size)
            run.font.color.rgb = rgb_from_hex(palette["paper"] if is_header else palette["ink"])

            if is_header or is_first_col:
                run.font.bold = True


def insert_png_before_table(table, png_path: Path, config: UCKKConfig) -> bool:
    try:
        paragraph = insert_paragraph_before_table(table)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, before=0, after=1)

        run = paragraph.add_run()
        run.add_picture(str(png_path), width=Pt(config.max_separator_width_pt))
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Helpers OOXML
# ---------------------------------------------------------------------------

def insert_paragraph_before_table(table) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._element.addprevious(new_p)
    return Paragraph(new_p, table._parent)


def set_table_layout_fixed(table) -> None:
    tbl_pr = get_or_add_tbl_pr(table)
    remove_children_by_tag(tbl_pr, "w:tblLayout")

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def set_table_borders(table, color_hex: str, width_pt: float) -> None:
    tbl_pr = get_or_add_tbl_pr(table)
    remove_children_by_tag(tbl_pr, "w:tblBorders")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(max(2, int(float(width_pt) * 8))))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), clean_hex(color_hex))
        borders.append(element)

    tbl_pr.append(borders)


def get_or_add_tbl_pr(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    return tbl_pr


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    remove_children_by_tag(tc_pr, "w:shd")

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), clean_hex(fill_hex))
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex: str, width_pt: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    remove_children_by_tag(tc_pr, "w:tcBorders")

    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(max(2, int(float(width_pt) * 8))))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), clean_hex(color_hex))
        borders.append(element)

    tc_pr.append(borders)


def set_cell_margins(cell, top: int = 18, bottom: int = 18, start: int = 58, end: int = 58) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    remove_children_by_tag(tc_pr, "w:tcMar")

    tc_mar = OxmlElement("w:tcMar")
    values = {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }

    for side, value in values.items():
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
        tc_mar.append(element)

    tc_pr.append(tc_mar)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()

    if not any(child.tag == qn("w:tblHeader") for child in tr_pr):
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()

    if not any(child.tag == qn("w:cantSplit") for child in tr_pr):
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def remove_children_by_tag(parent, tag: str) -> None:
    qualified = qn(tag)
    for child in list(parent):
        if child.tag == qualified:
            parent.remove(child)


def clean_hex(value: str) -> str:
    return str(value).replace("#", "").upper()


def rgb_from_hex(value: str) -> RGBColor:
    value = clean_hex(value)
    return RGBColor(
        int(value[0:2], 16),
        int(value[2:4], 16),
        int(value[4:6], 16),
    )


def parse_widths(value: str) -> List[float]:
    widths = []
    for part in str(value or "").split(","):
        part = part.strip()
        if not part:
            continue
        try:
            number = float(part)
            if number > 0:
                widths.append(number)
        except ValueError:
            continue
    return widths


def get_palette(mode: str) -> dict:
    if mode == "COULEUR":
        return {
            "petrol": "#123B42",
            "paper": "#FFFFFF",
            "surface": "#FFFFFF",
            "line": "#8DA0A6",
            "alt": "#F2F5F7",
            "left_col": "#DFE9EC",
            "ink": "#171717",
        }

    return {
        "petrol": "#2E3133",
        "paper": "#FFFFFF",
        "surface": "#FFFFFF",
        "line": "#A7A7A7",
        "alt": "#F1F1F1",
        "left_col": "#E4E4E4",
        "ink": "#111111",
    }


def log_msg(log: Optional[Callable[[str], None]], message: str) -> None:
    if log:
        log(message)


# ---------------------------------------------------------------------------
# Interface graphique Tkinter
# ---------------------------------------------------------------------------

class UCKKApp(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("UCKK — Stylisation simple des tableaux DOCX")
        self.geometry("760x650")
        self.minsize(680, 560)

        self._build_vars()
        self._build_ui()

        if IMPORT_ERROR is not None:
            messagebox.showerror(
                "Dépendance manquante",
                "Le module python-docx n'est pas installé.\n\n"
                "Installe-le avec :\n"
                "pip install python-docx",
            )

    def _build_vars(self) -> None:
        self.input_docx_var = tk.StringVar()
        self.output_docx_var = tk.StringVar()
        self.png_file_var = tk.StringVar()

        self.add_png_var = tk.BooleanVar(value=False)
        self.color_mode_var = tk.StringVar(value="BW")

        self.max_sep_width_var = tk.StringVar(value="310")
        self.border_width_var = tk.StringVar(value="0.5")
        self.header_font_size_var = tk.StringVar(value="8")
        self.body_font_size_var = tk.StringVar(value="8")
        self.font_title_var = tk.StringVar(value="Georgia")
        self.font_body_var = tk.StringVar(value="Arial")

        self.first_row_as_header_var = tk.BooleanVar(value=True)
        self.accent_first_col_var = tk.BooleanVar(value=True)
        self.alternate_rows_var = tk.BooleanVar(value=True)
        self.apply_widths_var = tk.BooleanVar(value=True)
        self.column_widths_var = tk.StringVar(value="120,175,175,175,175,175,175,175")

        self.skip_small_tables_var = tk.BooleanVar(value=False)
        self.min_rows_var = tk.StringVar(value="2")
        self.min_cols_var = tk.StringVar(value="2")

    def _build_ui(self) -> None:
        root = ttk.Frame(self, padding=12)
        root.pack(fill="both", expand=True)

        title = ttk.Label(root, text="UCKK — Tableaux DOCX", font=("Georgia", 20, "bold"))
        title.pack(anchor="w")

        subtitle = ttk.Label(
            root,
            text="Version simple : un seul PNG optionnel, aucun aperçu, stylisation directe.",
            foreground="#444444",
        )
        subtitle.pack(anchor="w", pady=(0, 10))

        notebook = ttk.Notebook(root)
        notebook.pack(fill="both", expand=True)

        tab_files = ttk.Frame(notebook, padding=10)
        tab_style = ttk.Frame(notebook, padding=10)
        tab_log = ttk.Frame(notebook, padding=10)

        notebook.add(tab_files, text="Fichiers")
        notebook.add(tab_style, text="Style")
        notebook.add(tab_log, text="Journal")

        self._build_files_tab(tab_files)
        self._build_style_tab(tab_style)
        self._build_log_tab(tab_log)

        bottom = ttk.Frame(root)
        bottom.pack(fill="x", pady=(10, 0))

        self.progress = ttk.Progressbar(bottom, mode="indeterminate")
        self.progress.pack(side="left", fill="x", expand=True, padx=(0, 8))

        self.run_button = ttk.Button(bottom, text="Styliser le DOCX", command=self.run_processing)
        self.run_button.pack(side="right")

    def _build_files_tab(self, parent) -> None:
        self._row_file(parent, "Fichier .docx source", self.input_docx_var, self.choose_input_docx)
        self._row_file(parent, "Fichier .docx de sortie", self.output_docx_var, self.choose_output_docx)
        self._row_file(parent, "PNG décoratif unique", self.png_file_var, self.choose_png_file)

        png_box = ttk.LabelFrame(parent, text="Option PNG", padding=10)
        png_box.pack(fill="x", pady=(12, 0))

        ttk.Checkbutton(
            png_box,
            text="Ajouter ce PNG avant chaque tableau",
            variable=self.add_png_var,
        ).pack(anchor="w")

        ttk.Button(png_box, text="Retirer le PNG", command=self.clear_png_file).pack(anchor="w", pady=(8, 0))

        hint = ttk.Label(
            parent,
            text="Si l'option PNG est décochée, l'app formate seulement les tableaux.",
            foreground="#555555",
        )
        hint.pack(anchor="w", pady=(12, 0))

        hint2 = ttk.Label(
            parent,
            text="Le fichier original n'est jamais modifié. Une copie .docx est créée.",
            foreground="#555555",
        )
        hint2.pack(anchor="w", pady=(4, 0))

    def _build_style_tab(self, parent) -> None:
        main = ttk.LabelFrame(parent, text="Style du tableau", padding=10)
        main.pack(fill="x")

        ttk.Label(main, text="Palette").grid(row=0, column=0, sticky="w")
        ttk.Combobox(
            main,
            textvariable=self.color_mode_var,
            values=["BW", "COULEUR"],
            state="readonly",
            width=18,
        ).grid(row=0, column=1, sticky="w", padx=(8, 0))

        checks = ttk.LabelFrame(parent, text="Options", padding=10)
        checks.pack(fill="x", pady=(12, 0))

        for text, var in [
            ("Première ligne = en-tête", self.first_row_as_header_var),
            ("Accentuer la première colonne", self.accent_first_col_var),
            ("Alterner légèrement les lignes", self.alternate_rows_var),
            ("Appliquer les largeurs de colonnes", self.apply_widths_var),
            ("Ignorer les petits tableaux", self.skip_small_tables_var),
        ]:
            ttk.Checkbutton(checks, text=text, variable=var).pack(anchor="w", pady=2)

        measures = ttk.LabelFrame(parent, text="Mesures", padding=10)
        measures.pack(fill="x", pady=(12, 0))

        self._grid_entry(measures, 0, "Largeur PNG max en points", self.max_sep_width_var)
        self._grid_entry(measures, 1, "Épaisseur bordure", self.border_width_var)
        self._grid_entry(measures, 2, "Taille police en-tête", self.header_font_size_var)
        self._grid_entry(measures, 3, "Taille police corps", self.body_font_size_var)
        self._grid_entry(measures, 4, "Nombre min. de lignes", self.min_rows_var)
        self._grid_entry(measures, 5, "Nombre min. de colonnes", self.min_cols_var)

        fonts = ttk.LabelFrame(parent, text="Polices et colonnes", padding=10)
        fonts.pack(fill="x", pady=(12, 0))

        self._grid_entry(fonts, 0, "Police titres/en-têtes", self.font_title_var)
        self._grid_entry(fonts, 1, "Police corps", self.font_body_var)

        ttk.Label(fonts, text="Largeurs colonnes en points").grid(row=2, column=0, sticky="w", pady=3)
        ttk.Entry(fonts, textvariable=self.column_widths_var, width=55).grid(
            row=2, column=1, sticky="ew", padx=(8, 0), pady=3
        )
        fonts.columnconfigure(1, weight=1)

    def _build_log_tab(self, parent) -> None:
        self.log_text = tk.Text(parent, height=24, wrap="word")
        self.log_text.pack(fill="both", expand=True)

        buttons = ttk.Frame(parent)
        buttons.pack(fill="x", pady=(8, 0))
        ttk.Button(buttons, text="Effacer le journal", command=lambda: self.log_text.delete("1.0", "end")).pack(side="left")

        self.log("Prêt.")

    def _row_file(self, parent, label: str, variable: tk.StringVar, command) -> None:
        frame = ttk.Frame(parent)
        frame.pack(fill="x", pady=5)

        ttk.Label(frame, text=label, width=24).pack(side="left")
        ttk.Entry(frame, textvariable=variable).pack(side="left", fill="x", expand=True, padx=(8, 8))
        ttk.Button(frame, text="Choisir", command=command).pack(side="right")

    def _grid_entry(self, parent, row: int, label: str, variable: tk.StringVar) -> None:
        ttk.Label(parent, text=label).grid(row=row, column=0, sticky="w", pady=3)
        ttk.Entry(parent, textvariable=variable, width=24).grid(row=row, column=1, sticky="w", padx=(8, 0), pady=3)

    def choose_input_docx(self) -> None:
        path = filedialog.askopenfilename(
            title="Choisir le fichier DOCX source",
            filetypes=[("Word document", "*.docx")],
        )
        if not path:
            return

        self.input_docx_var.set(path)

        if not self.output_docx_var.get().strip():
            src = Path(path)
            self.output_docx_var.set(str(src.with_name(src.stem + "_UCKK_tableaux.docx")))

    def choose_output_docx(self) -> None:
        path = filedialog.asksaveasfilename(
            title="Choisir le fichier DOCX de sortie",
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
        )
        if path:
            self.output_docx_var.set(path)

    def choose_png_file(self) -> None:
        path = filedialog.askopenfilename(
            title="Choisir un PNG décoratif",
            filetypes=[("PNG", "*.png")],
        )
        if path:
            self.png_file_var.set(path)
            self.add_png_var.set(True)

    def clear_png_file(self) -> None:
        self.png_file_var.set("")
        self.add_png_var.set(False)

    def gather_config(self) -> UCKKConfig:
        return UCKKConfig(
            input_docx=self.input_docx_var.get().strip(),
            output_docx=self.output_docx_var.get().strip(),
            png_file=self.png_file_var.get().strip(),
            add_png=self.add_png_var.get(),
            color_mode=self.color_mode_var.get(),
            max_separator_width_pt=float_or_default(self.max_sep_width_var.get(), 310.0),
            table_border_width_pt=float_or_default(self.border_width_var.get(), 0.5),
            header_font_size=float_or_default(self.header_font_size_var.get(), 8.0),
            body_font_size=float_or_default(self.body_font_size_var.get(), 8.0),
            font_title=self.font_title_var.get().strip() or "Georgia",
            font_body=self.font_body_var.get().strip() or "Arial",
            first_row_as_header=self.first_row_as_header_var.get(),
            accent_first_column=self.accent_first_col_var.get(),
            alternate_rows=self.alternate_rows_var.get(),
            apply_column_widths=self.apply_widths_var.get(),
            column_widths_pt=self.column_widths_var.get().strip(),
            skip_small_tables=self.skip_small_tables_var.get(),
            min_rows=int_or_default(self.min_rows_var.get(), 2),
            min_cols=int_or_default(self.min_cols_var.get(), 2),
        )

    def run_processing(self) -> None:
        try:
            cfg = self.gather_config()

            if not cfg.input_docx:
                raise ValueError("Choisis un fichier .docx source.")

            if not cfg.output_docx:
                src = Path(cfg.input_docx)
                cfg.output_docx = str(src.with_name(src.stem + "_UCKK_tableaux.docx"))
                self.output_docx_var.set(cfg.output_docx)

            # Valide tôt le PNG si l'option est activée.
            resolve_png_path(cfg)

        except Exception as exc:
            messagebox.showerror("Réglages incomplets", str(exc))
            return

        self.run_button.configure(state="disabled")
        self.progress.start(10)
        self.log("\nLancement de la stylisation…")

        thread = threading.Thread(target=self._worker, args=(cfg,), daemon=True)
        thread.start()

    def _worker(self, cfg: UCKKConfig) -> None:
        try:
            result = process_docx(cfg, log=lambda msg: self.after(0, self.log, msg))
            self.after(0, self._finish_success, cfg, result)
        except Exception as exc:
            tb = traceback.format_exc()
            self.after(0, self._finish_error, exc, tb)

    def _finish_success(self, cfg: UCKKConfig, result: ProcessingResult) -> None:
        self.progress.stop()
        self.run_button.configure(state="normal")

        msg = (
            "Terminé.\n\n"
            f"Fichier créé : {cfg.output_docx}\n\n"
            f"Tableaux trouvés : {result.tables_found}\n"
            f"Tableaux stylisés : {result.tables_styled}\n"
            f"Tableaux ignorés : {result.tables_skipped}\n"
            f"Séparateurs insérés : {result.separators_inserted}"
        )

        if result.warnings:
            msg += "\n\nAvertissements :\n- " + "\n- ".join(result.warnings)

        self.log("\n" + msg)
        messagebox.showinfo("Stylisation terminée", msg)

    def _finish_error(self, exc: Exception, tb: str) -> None:
        self.progress.stop()
        self.run_button.configure(state="normal")

        self.log("\nERREUR\n" + tb)
        messagebox.showerror("Erreur", str(exc))

    def log(self, message: str) -> None:
        self.log_text.insert("end", str(message) + "\n")
        self.log_text.see("end")


def float_or_default(value: str, fallback: float) -> float:
    try:
        return float(str(value).replace(",", "."))
    except Exception:
        return fallback


def int_or_default(value: str, fallback: int) -> int:
    try:
        return int(float(str(value).replace(",", ".")))
    except Exception:
        return fallback


def main() -> None:
    app = UCKKApp()
    app.mainloop()


if __name__ == "__main__":
    main()
