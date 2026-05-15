#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
king_klown_premium_builder.py

Moteur de standardisation premium pour manuels King Klown / UCKK.

Entrée :
    .docx exporté de Google Docs ou Word

Sorties :
    .docx standardisé
    .pdf optionnel si LibreOffice est installé

Dépendance :
    pip install python-docx

Lancer le GUI :
    python king_klown_premium_builder.py

CLI :
    python king_klown_premium_builder.py "manuel.docx" --toc
    python king_klown_premium_builder.py "./docs" --recursive --margin-mode recto-verso --pdf

Notes :
- Le script garde seulement le layout Premium.
- Couleur de marque unique : #1E6864.
- Il corrige les paragraphes courts justifiés qui créent des espaces énormes.
- Par défaut, le corps est aligné à gauche pour éviter les blancs étirés.
- Il ignore les titres génériques comme "Word Document" dans le header.
- Le header peut suivre automatiquement le chapitre courant via STYLEREF.
- Il ajoute un mode marge centrée, marge gauche ou recto-verso.
"""

from __future__ import annotations

import argparse
import queue
import re
import shutil
import subprocess
import sys
import threading
import time
import tkinter as tk
from dataclasses import dataclass, field
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
from typing import Callable, Iterable, Optional
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from docx import Document
    from docx.document import Document as DocxDocument
    from docx.enum.section import WD_SECTION_START
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError as exc:
    raise SystemExit(
        "Dépendance manquante : python-docx.\n"
        "Installe-la avec : pip install python-docx"
    ) from exc


# ---------------------------------------------------------------------------
# Charte Premium King Klown / UCKK
# ---------------------------------------------------------------------------

UCKK_DARK = "1B1B1B"
UCKK_SOFT_BLACK = "2A2A2A"
UCKK_GRAY = "666666"
UCKK_LIGHT_GRAY = "F4F4F2"
UCKK_BORDER = "D7D2C8"
UCKK_GOLD = "1E6864"  # accent principal
UCKK_DEEP_RED = "1E6864"  # couleur de marque King Klown / UCKK
UCKK_PARCHMENT = "E8F3F1"
UCKK_PALE_GOLD = "E8F3F1"
WHITE = "FFFFFF"


@dataclass
class PremiumConfig:
    page_size: str = "letter"             # letter ou a4
    margin_mode: str = "centered"         # centered, left, recto-verso

    # Layout premium par défaut.
    margin_top_cm: float = 2.2
    margin_bottom_cm: float = 2.0
    margin_left_cm: float = 2.25
    margin_right_cm: float = 2.25
    margin_inside_cm: float = 2.75
    margin_outside_cm: float = 1.95

    body_font: str = "EB Garamond"
    heading_font: str = "EB Garamond"
    mono_font: str = "Consolas"

    body_size_pt: float = 10.7
    line_spacing: float = 1.045

    add_toc: bool = False
    add_header_footer: bool = True
    header_layout: str = "chapter"         # chapter, brand-chapter, centered, split, minimal, none
    body_alignment: str = "left"          # left, smart-justify, justify
    footer_text: str = "Univers-Cité King Klown"

    chapter_page_breaks: bool = True
    clear_google_run_formatting: bool = True
    normalize_heading_levels: bool = True
    standardize_tables: bool = True
    add_callout_boxes: bool = True
    export_pdf: bool = False


@dataclass
class ProcessStats:
    source: Path
    output: Path
    headings: int = 0
    callouts: int = 0
    formula_blocks: int = 0
    code_lines: int = 0
    tables: int = 0
    blank_paragraphs_removed: int = 0
    warnings: list[str] = field(default_factory=list)
    dry_run: bool = False


# ---------------------------------------------------------------------------
# OOXML helpers
# ---------------------------------------------------------------------------

def _hex(rgb: str) -> str:
    return rgb.replace("#", "").upper()


def get_or_add(parent, child_tag: str):
    child = parent.find(qn(child_tag))
    if child is None:
        child = OxmlElement(child_tag)
        parent.append(child)
    return child


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = get_or_add(tc_pr, "w:shd")
    shd.set(qn("w:fill"), _hex(fill))


def set_cell_borders(cell, color: str = UCKK_BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = get_or_add(tc_pr, "w:tcBorders")

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), _hex(color))


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = get_or_add(tc_pr, "w:tcMar")
    values = {"top": top, "start": start, "bottom": bottom, "end": end}

    for side, value in values.items():
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = get_or_add(p_pr, "w:shd")
    shd.set(qn("w:fill"), _hex(fill))


def set_paragraph_border(paragraph, side: str = "left", color: str = UCKK_GOLD, size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = get_or_add(p_pr, "w:pBdr")

    border = borders.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        borders.append(border)

    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "6")
    border.set(qn("w:color"), _hex(color))


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def add_field(paragraph, instr: str, placeholder: str = "") -> None:
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    run._r.append(instr_text)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        run._r.append(t)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def enable_update_fields_on_open(docx_path: Path) -> None:
    tmp_path = docx_path.with_suffix(".tmp.docx")
    with ZipFile(docx_path, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                xml = data.decode("utf-8")
                if "w:updateFields" not in xml:
                    xml = xml.replace("</w:settings>", '<w:updateFields w:val="true"/></w:settings>')
                    data = xml.encode("utf-8")
            zout.writestr(item, data)
    tmp_path.replace(docx_path)


def enable_mirror_margins(doc: DocxDocument) -> None:
    settings = doc.settings._element
    if settings.find(qn("w:mirrorMargins")) is None:
        settings.append(OxmlElement("w:mirrorMargins"))


def clear_run_formatting_keep_emphasis(run) -> None:
    r_pr = run._r.rPr
    if r_pr is None:
        return

    # On supprime les surcharges qui viennent souvent de Google Docs,
    # mais on garde gras, italique, souligné, exposant, indice.
    for tag in [
        "w:rFonts", "w:sz", "w:szCs", "w:color", "w:highlight", "w:shd",
        "w:spacing", "w:kern", "w:position", "w:lang",
    ]:
        for child in list(r_pr.findall(qn(tag))):
            r_pr.remove(child)


# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------

def style_lookup(doc: DocxDocument, *names: str):
    lower = {s.name.lower(): s for s in doc.styles if getattr(s, "name", None)}
    for name in names:
        try:
            return doc.styles[name]
        except KeyError:
            pass
        found = lower.get(name.lower())
        if found is not None:
            return found
    return None


def ensure_paragraph_style(doc: DocxDocument, name: str, base: Optional[str] = None):
    style = style_lookup(doc, name)
    if style is not None:
        return style

    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base:
        base_style = style_lookup(doc, base)
        if base_style is not None:
            style.base_style = base_style
    return style


def set_font(style, name: str, size_pt: Optional[float] = None,
             bold: Optional[bool] = None, italic: Optional[bool] = None,
             color: Optional[str] = None) -> None:
    style.font.name = name
    if style._element.rPr is not None and style._element.rPr.rFonts is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        style._element.rPr.rFonts.set(qn("w:cs"), name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(_hex(color))


def alignment_from_body_setting(body_alignment: str):
    """Retourne l'alignement Word pour le corps.

    left est le défaut premium, parce que les exports Google Docs/Word
    contiennent souvent des retours de ligne manuels. En justification,
    Word étire alors chaque ligne courte et crée des blancs énormes.
    """
    if body_alignment == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if body_alignment == "smart-justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def configure_styles(doc: DocxDocument, cfg: PremiumConfig) -> None:
    normal = style_lookup(doc, "Normal", "normal") or ensure_paragraph_style(doc, "Normal")
    set_font(normal, cfg.body_font, cfg.body_size_pt, color=UCKK_DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5.2)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = cfg.line_spacing
    normal.paragraph_format.alignment = alignment_from_body_setting(cfg.body_alignment)

    h1 = style_lookup(doc, "Heading 1", "Titre 1") or ensure_paragraph_style(doc, "Heading 1", "Normal")
    set_font(h1, cfg.heading_font, 17.5, bold=True, color=UCKK_DEEP_RED)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(9)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h2 = style_lookup(doc, "Heading 2", "Titre 2") or ensure_paragraph_style(doc, "Heading 2", "Normal")
    set_font(h2, cfg.heading_font, 13.2, bold=True, color=UCKK_SOFT_BLACK)
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h3 = style_lookup(doc, "Heading 3", "Titre 3") or ensure_paragraph_style(doc, "Heading 3", "Normal")
    set_font(h3, cfg.heading_font, 11.6, bold=True, color=UCKK_SOFT_BLACK)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(3.5)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h4 = style_lookup(doc, "Heading 4", "Titre 4") or ensure_paragraph_style(doc, "Heading 4", "Normal")
    set_font(h4, cfg.heading_font, 10.6, bold=True, color=UCKK_GRAY)
    h4.paragraph_format.space_before = Pt(9)
    h4.paragraph_format.space_after = Pt(2)
    h4.paragraph_format.keep_with_next = True
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = ensure_paragraph_style(doc, "KK - Intro", "Normal")
    set_font(intro, cfg.body_font, cfg.body_size_pt + 0.5, color=UCKK_DARK)
    intro.paragraph_format.space_before = Pt(4)
    intro.paragraph_format.space_after = Pt(9)
    intro.paragraph_format.line_spacing = 1.08
    intro.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    callout = ensure_paragraph_style(doc, "KK - Encadre", "Normal")
    set_font(callout, cfg.body_font, cfg.body_size_pt, color=UCKK_DARK)
    callout.paragraph_format.left_indent = Cm(0.38)
    callout.paragraph_format.right_indent = Cm(0.25)
    callout.paragraph_format.space_before = Pt(7)
    callout.paragraph_format.space_after = Pt(7)
    callout.paragraph_format.line_spacing = 1.04
    callout.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    formula = ensure_paragraph_style(doc, "KK - Formule", "Normal")
    set_font(formula, cfg.body_font, cfg.body_size_pt + 0.2, italic=False, color=UCKK_DARK)
    formula.paragraph_format.left_indent = Cm(0.4)
    formula.paragraph_format.space_before = Pt(5)
    formula.paragraph_format.space_after = Pt(5)
    formula.paragraph_format.line_spacing = 1.05
    formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    transformation = ensure_paragraph_style(doc, "KK - Transformation", "Normal")
    set_font(transformation, cfg.body_font, cfg.body_size_pt, italic=True, color=UCKK_GOLD)
    transformation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    transformation.paragraph_format.space_before = Pt(6)
    transformation.paragraph_format.space_after = Pt(6)
    transformation.paragraph_format.line_spacing = 1.0

    code = ensure_paragraph_style(doc, "KK - Code", "Normal")
    set_font(code, cfg.mono_font, 9.2, color=UCKK_DARK)
    code.paragraph_format.left_indent = Cm(0.35)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    caption = ensure_paragraph_style(doc, "KK - Legende", "Normal")
    set_font(caption, cfg.body_font, 9.6, italic=True, color=UCKK_GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Détection de structure
# ---------------------------------------------------------------------------

H4_RE = re.compile(r"^\s*\d{1,2}\.\d+\.\d+\.\d+\s+.+")
H3_RE = re.compile(r"^\s*\d{1,2}\.\d+\.\d+\s+.+")
H2_RE = re.compile(r"^\s*\d{1,2}\.\d+\s+.+")
H1_RE = re.compile(r"^\s*\d{2}\.\s+.+")
PART_RE = re.compile(r"^\s*Partie\s+\d+\s+[—-]\s+.+", re.IGNORECASE)
CHAPTER_RE = re.compile(r"^\s*(Chapitre|Chapter)\s+\d+\s+[—-]\s+.+", re.IGNORECASE)

CALLOUT_RE = re.compile(
    r"^\s*(Phrase clé|Point de vigilance|Formule courte|Formule officielle|"
    r"Devise|Règle|À retenir|Retenir|Erreur fréquente|Erreurs fréquentes|"
    r"Production attendue|Critères de réussite|Critères|Trace à conserver|"
    r"Livrable|Mini-exercice|Exercice|Objectifs du chapitre|Attention|Vigilance)\b",
    re.IGNORECASE,
)

GENERIC_TITLES = {
    "word document",
    "document",
    "untitled document",
    "sans titre",
    "document sans titre",
}


def clean_markdown_artifacts(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = re.sub(r"^\*\*(.+)\*\*$", r"\1", text)
    text = re.sub(r"^__(.+)__$", r"\1", text)
    text = text.replace("\\#", "#")
    text = text.replace("\\_", "_")
    text = text.replace("\\*", "*")
    return text.strip()


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        for i, run in enumerate(paragraph.runs):
            run.text = new_text if i == 0 else ""
    else:
        paragraph.add_run(new_text)


def strip_markdown_heading_marks(paragraph) -> None:
    cleaned = clean_markdown_artifacts(paragraph.text)
    replace_paragraph_text(paragraph, cleaned)


def classify_heading(text: str) -> Optional[str]:
    stripped = text.strip()
    if not stripped:
        return None

    if stripped.startswith("#### "):
        return "Heading 4"
    if stripped.startswith("### "):
        return "Heading 3"
    if stripped.startswith("## "):
        return "Heading 2"
    if stripped.startswith("# "):
        return "Heading 1"

    cleaned = clean_markdown_artifacts(stripped)

    if H4_RE.match(cleaned):
        return "Heading 4"
    if H3_RE.match(cleaned):
        return "Heading 3"
    if H2_RE.match(cleaned):
        return "Heading 2"
    if H1_RE.match(cleaned) or PART_RE.match(cleaned) or CHAPTER_RE.match(cleaned):
        return "Heading 1"
    return None


def is_list_paragraph(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.numPr is not None


def has_manual_line_breaks(paragraph) -> bool:
    """Google Docs et Word exportent parfois des retours de ligne manuels.
    En texte justifié, ces lignes sont étirées. On les repère pour les mettre
    à gauche ou en formule.
    """
    if "\n" in paragraph.text:
        return True
    return any("\n" in run.text for run in paragraph.runs)


def is_spread_risk_text(text: str) -> bool:
    cleaned = clean_markdown_artifacts(text)
    if not cleaned:
        return False

    lines = [ln.strip() for ln in cleaned.splitlines() if ln.strip()]
    if len(lines) > 1:
        # Plusieurs lignes courtes dans un même paragraphe = grand risque
        # d'espaces étirés en justification.
        return any(len(ln) <= 120 for ln in lines)

    if len(cleaned) <= 140:
        return True

    # Les phrases avec peu de mots longs et beaucoup d'espaces potentiels
    # se déforment facilement dans un bloc justifié.
    words = cleaned.split()
    return len(words) <= 12



def looks_like_formula_block(text: str) -> bool:
    """Détecte les formules courtes qui ne doivent jamais être justifiées.

    Corrige notamment le problème visible dans Word :
    "Inventorier      avant      d'analyser"
    causé par un paragraphe justifié avec retours de ligne manuels.
    """
    cleaned = clean_markdown_artifacts(text)
    lines = [ln.strip() for ln in cleaned.splitlines() if ln.strip()]

    if not lines:
        return False

    if "→" in cleaned and len(cleaned) <= 220:
        return True

    if len(lines) >= 2 and all(len(ln) <= 90 for ln in lines):
        formula_hits = sum(
            1 for ln in lines
            if re.search(r"\bavant\b|\b→\b|:$", ln, flags=re.IGNORECASE)
        )
        return formula_hits >= 1

    if len(cleaned) <= 110 and re.search(r"\bavant\b", cleaned, flags=re.IGNORECASE):
        return True

    return False


def looks_like_code_line(text: str) -> bool:
    cleaned = clean_markdown_artifacts(text)
    return bool(re.match(r"^\s*(```|[a-zA-Z0-9_]+\s*=\s*|[A-Za-z0-9_./-]+\.(py|md|docx|pdf))", cleaned))


def standardize_paragraphs(doc: DocxDocument, cfg: PremiumConfig, stats: ProcessStats) -> None:
    first_h1_seen = False
    normal = style_lookup(doc, "Normal", "normal")
    formula_style = style_lookup(doc, "KK - Formule")
    transformation_style = style_lookup(doc, "KK - Transformation")
    code_style = style_lookup(doc, "KK - Code")
    callout_style = style_lookup(doc, "KK - Encadre")

    for p in doc.paragraphs:
        text = p.text.strip()

        if cfg.clear_google_run_formatting:
            for run in p.runs:
                clear_run_formatting_keep_emphasis(run)

        if cfg.normalize_heading_levels:
            level = classify_heading(text)
            if level:
                strip_markdown_heading_marks(p)
                target = style_lookup(doc, level, f"Titre {level[-1]}")
                if target is not None:
                    p.style = target
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                if level == "Heading 1":
                    if cfg.chapter_page_breaks and first_h1_seen:
                        p.paragraph_format.page_break_before = True
                    first_h1_seen = True

                stats.headings += 1
                continue

        # Listes : jamais justifiées.
        if is_list_paragraph(p):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.03
            p.paragraph_format.space_after = Pt(2.5)
            continue

        if text and CALLOUT_RE.match(clean_markdown_artifacts(text)):
            if callout_style is not None:
                p.style = callout_style
            if cfg.add_callout_boxes:
                set_paragraph_shading(p, UCKK_PARCHMENT)
                set_paragraph_border(p, "left", UCKK_GOLD, "14")
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats.callouts += 1
            continue

        if text and looks_like_formula_block(text):
            if "→" in text and transformation_style is not None:
                p.style = transformation_style
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif formula_style is not None:
                p.style = formula_style
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats.formula_blocks += 1
            continue

        if text and looks_like_code_line(text):
            if code_style is not None:
                p.style = code_style
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stats.code_lines += 1
            continue

        if normal is not None and p.style is not None:
            if p.style.name.lower() in {"normal", "normal text", "body text"}:
                p.style = normal

        # Anti-spread : les paragraphes courts, les lignes manuelles et les
        # blocs semi-formulaires ne doivent jamais être justifiés.
        if text:
            if cfg.body_alignment == "left":
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif cfg.body_alignment == "smart-justify":
                if has_manual_line_breaks(p) or is_spread_risk_text(text):
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif cfg.body_alignment == "justify":
                if has_manual_line_breaks(p) and is_spread_risk_text(text):
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def remove_excess_blank_paragraphs(doc: DocxDocument, stats: ProcessStats, max_blanks: int = 1) -> None:
    blanks = 0
    for p in list(doc.paragraphs):
        if p.text.strip():
            blanks = 0
            continue
        blanks += 1
        if blanks > max_blanks:
            remove_paragraph(p)
            stats.blank_paragraphs_removed += 1


# ---------------------------------------------------------------------------
# Page, header, footer, tables, TOC
# ---------------------------------------------------------------------------

def configure_sections(doc: DocxDocument, cfg: PremiumConfig) -> None:
    if cfg.margin_mode == "recto-verso":
        enable_mirror_margins(doc)
        left = cfg.margin_inside_cm
        right = cfg.margin_outside_cm
    elif cfg.margin_mode == "left":
        left = 2.75
        right = 1.95
    else:
        left = cfg.margin_left_cm
        right = cfg.margin_right_cm

    for section in doc.sections:
        if cfg.page_size.lower() == "a4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        else:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)

        section.top_margin = Cm(cfg.margin_top_cm)
        section.bottom_margin = Cm(cfg.margin_bottom_cm)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(0.85)
        section.start_type = WD_SECTION_START.NEW_PAGE


def get_document_title(doc: DocxDocument, fallback: str = "Manuel King Klown") -> str:
    core_title = (doc.core_properties.title or "").strip()
    if core_title and core_title.lower() not in GENERIC_TITLES:
        return core_title[:90]

    for p in doc.paragraphs:
        t = clean_markdown_artifacts(p.text)
        if not t:
            continue
        if t.lower() in GENERIC_TITLES:
            continue
        if len(t) > 90:
            return t[:90].rstrip() + "…"
        return t

    return fallback


def get_heading_1_style_name(doc: DocxDocument) -> str:
    """Retourne le nom réel du style Heading 1 / Titre 1.

    Le champ Word STYLEREF a besoin du nom du style pour afficher
    automatiquement le titre du chapitre courant dans le header.
    """
    style = style_lookup(doc, "Heading 1", "Titre 1")
    if style is not None:
        return style.name
    return "Heading 1"


def add_current_chapter_field(paragraph, doc: DocxDocument, placeholder: str = "Chapitre courant") -> None:
    """Insère un champ Word dynamique qui affiche le dernier Heading 1 actif.

    Dans Word/LibreOffice, le champ STYLEREF se met à jour à l'ouverture
    ou lors de l'export PDF. Il permet d'avoir un header qui suit les chapitres
    sans créer manuellement une section par chapitre.
    """
    style_name = get_heading_1_style_name(doc)
    safe_style_name = style_name.replace('"', r'\"')
    add_field(paragraph, f'STYLEREF "{safe_style_name}" \\* MERGEFORMAT', placeholder)


def format_run(run, cfg: PremiumConfig, size: float = 8.5, color: str = UCKK_GRAY, bold: bool = False) -> None:
    run.font.name = cfg.heading_font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(_hex(color))
    run.font.bold = bold


def add_header_footer(doc: DocxDocument, cfg: PremiumConfig) -> None:
    title = get_document_title(doc)

    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = ""

        if cfg.header_layout == "none":
            pass

        elif cfg.header_layout == "chapter":
            # Recommandé : affiche automatiquement le Heading 1 courant.
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_current_chapter_field(p, doc, title)
            for run in p.runs:
                format_run(run, cfg, size=8.2, color=UCKK_GRAY, bold=False)

        elif cfg.header_layout == "brand-chapter":
            # Variante : marque + chapitre courant.
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            brand = p.add_run("Univers-Cité King Klown")
            format_run(brand, cfg, size=8.2, color=UCKK_DEEP_RED, bold=True)
            sep = p.add_run(" — ")
            format_run(sep, cfg, size=8.2, color=UCKK_GRAY, bold=False)
            add_current_chapter_field(p, doc, title)
            for run in p.runs:
                if run.font.size is None:
                    format_run(run, cfg, size=8.2, color=UCKK_GRAY, bold=False)

        elif cfg.header_layout == "split":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            left = p.add_run("Univers-Cité King Klown")
            format_run(left, cfg, size=8.2, color=UCKK_DEEP_RED, bold=True)
            p.add_run("	")
            add_current_chapter_field(p, doc, title)
            for run in p.runs:
                if run.font.size is None:
                    format_run(run, cfg, size=8.2, color=UCKK_GRAY, bold=False)

        elif cfg.header_layout == "minimal":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("King Klown")
            format_run(r, cfg, size=8.2, color=UCKK_GRAY, bold=True)

        else:
            # Header statique, surtout utile pour les documents courts.
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(title)
            format_run(r, cfg, size=8.4, color=UCKK_GRAY, bold=False)

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        left = fp.add_run(cfg.footer_text + " — Page ")
        format_run(left, cfg, size=8.2, color=UCKK_GRAY)

        add_field(fp, "PAGE", "1")
        fp.add_run(" / ")
        add_field(fp, "NUMPAGES", "1")

        for run in fp.runs:
            if run.font.size is None:
                format_run(run, cfg, size=8.2, color=UCKK_GRAY)


def insert_toc_at_start(doc: DocxDocument) -> None:
    if not doc.paragraphs:
        doc.add_paragraph()

    first = doc.paragraphs[0]

    for p in doc.paragraphs[:10]:
        if p.text.strip().lower() in {"table des matières", "table of contents", "sommaire"}:
            return

    toc_title = first.insert_paragraph_before("Table des matières", style="Heading 1")
    toc_title.paragraph_format.page_break_before = False

    toc_p = first.insert_paragraph_before("")
    add_field(toc_p, r'TOC \o "1-3" \h \z \u', "Cliquez ici puis actualisez la table des matières.")

    sep = first.insert_paragraph_before("")
    sep.add_run().add_break(WD_BREAK.PAGE)


def standardize_tables(doc: DocxDocument, cfg: PremiumConfig, stats: ProcessStats) -> None:
    normal = style_lookup(doc, "Normal", "normal")
    stats.tables = len(doc.tables)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_borders(cell)
                set_cell_margins(cell)

                if row_idx == 0:
                    set_cell_shading(cell, UCKK_DEEP_RED)
                elif row_idx % 2 == 0:
                    set_cell_shading(cell, UCKK_LIGHT_GRAY)
                else:
                    set_cell_shading(cell, WHITE)

                for p in cell.paragraphs:
                    if normal is not None:
                        p.style = normal
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    for run in p.runs:
                        if cfg.clear_google_run_formatting:
                            clear_run_formatting_keep_emphasis(run)
                        run.font.name = cfg.body_font
                        run.font.size = Pt(9.7)
                        if row_idx == 0:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor.from_string(WHITE)
                        else:
                            run.font.color.rgb = RGBColor.from_string(UCKK_DARK)


def mark_header_rows_repeat(doc: DocxDocument) -> None:
    for table in doc.tables:
        if not table.rows:
            continue
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = tr_pr.find(qn("w:tblHeader"))
        if tbl_header is None:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------

def backup_file(path: Path) -> Path:
    stamp = time.strftime("%Y%m%d_%H%M%S")
    backup = path.with_suffix(path.suffix + f".bak_{stamp}")
    backup.write_bytes(path.read_bytes())
    return backup


def validate_input_path(input_path: Path) -> None:
    if not input_path.exists():
        raise FileNotFoundError(f"Introuvable : {input_path}")
    if input_path.is_file() and input_path.suffix.lower() != ".docx":
        raise ValueError(f"Le fichier n'est pas un .docx : {input_path}")


def process_docx(input_path: Path, output_path: Path, cfg: PremiumConfig,
                 dry_run: bool = False, in_place: bool = False) -> ProcessStats:
    validate_input_path(input_path)

    stats = ProcessStats(source=input_path, output=output_path, dry_run=dry_run)
    doc = Document(str(input_path))

    configure_sections(doc, cfg)
    configure_styles(doc, cfg)
    standardize_paragraphs(doc, cfg, stats)
    remove_excess_blank_paragraphs(doc, stats)

    if cfg.standardize_tables:
        standardize_tables(doc, cfg, stats)
        mark_header_rows_repeat(doc)

    if cfg.add_header_footer:
        add_header_footer(doc, cfg)

    if cfg.add_toc:
        insert_toc_at_start(doc)

    if dry_run:
        return stats

    output_path.parent.mkdir(parents=True, exist_ok=True)

    if in_place:
        backup = backup_file(input_path)
        stats.warnings.append(f"Sauvegarde créée : {backup.name}")

    doc.save(str(output_path))

    if cfg.add_toc or cfg.header_layout in {"chapter", "brand-chapter", "split"}:
        try:
            enable_update_fields_on_open(output_path)
        except Exception as exc:
            stats.warnings.append(f"Champs Word : impossible d'activer updateFields : {exc}")

    if cfg.export_pdf:
        try:
            pdf_path = export_to_pdf(output_path)
            stats.warnings.append(f"PDF exporté : {pdf_path.name}")
        except Exception as exc:
            stats.warnings.append(f"Export PDF ignoré : {exc}")

    return stats


def iter_docx_files(path: Path, recursive: bool) -> Iterable[Path]:
    if path.is_file():
        if path.suffix.lower() == ".docx" and not path.name.startswith("~$"):
            yield path
        return

    pattern = "**/*.docx" if recursive else "*.docx"
    for p in sorted(path.glob(pattern)):
        if p.is_file() and not p.name.startswith("~$"):
            yield p


def output_for(input_file: Path, input_root: Path, output_root: Path, suffix: str, in_place: bool) -> Path:
    if in_place:
        return input_file

    if input_root.is_file():
        rel = Path(input_file.name)
    else:
        rel = input_file.relative_to(input_root)

    out = output_root / rel
    return out.with_name(out.stem + suffix + out.suffix)


def find_libreoffice() -> Optional[str]:
    for candidate in ("soffice", "libreoffice"):
        found = shutil.which(candidate)
        if found:
            return found

    win_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for p in win_paths:
        if Path(p).exists():
            return p

    return None


def export_to_pdf(docx_path: Path) -> Path:
    soffice = find_libreoffice()
    if not soffice:
        raise RuntimeError("LibreOffice introuvable.")

    out_dir = docx_path.parent
    cmd = [
        soffice, "--headless", "--convert-to", "pdf",
        "--outdir", str(out_dir), str(docx_path),
    ]
    completed = subprocess.run(cmd, capture_output=True, text=True, check=False)
    if completed.returncode != 0:
        raise RuntimeError(completed.stderr.strip() or completed.stdout.strip() or "Échec LibreOffice")

    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError("PDF attendu non créé.")
    return pdf_path


def run_batch(input_root: Path, output_root: Path, cfg: PremiumConfig,
              recursive: bool = False, in_place: bool = False,
              suffix: str = "_KK_PREMIUM", dry_run: bool = False,
              logger: Optional[Callable[[str], None]] = None) -> tuple[list[ProcessStats], int]:
    log = logger or print
    validate_input_path(input_root)

    files = list(iter_docx_files(input_root, recursive))
    if not files:
        log("Aucun fichier .docx trouvé.")
        return [], 0

    log(f"{len(files)} fichier(s) à traiter.")
    stats_list: list[ProcessStats] = []
    errors = 0

    for input_file in files:
        out = output_for(input_file, input_root, output_root, suffix, in_place)
        try:
            stats = process_docx(input_file, out, cfg, dry_run=dry_run, in_place=in_place)
            stats_list.append(stats)
            prefix = "DRY-RUN" if dry_run else "OK"
            log(
                f"{prefix}  {input_file.name} → {out}\n"
                f"      titres={stats.headings}, encadrés={stats.callouts}, "
                f"formules={stats.formula_blocks}, code={stats.code_lines}, "
                f"tableaux={stats.tables}, blancs_supprimés={stats.blank_paragraphs_removed}"
            )
            for warning in stats.warnings:
                log(f"      note: {warning}")
        except Exception as exc:
            errors += 1
            log(f"ERREUR  {input_file} : {exc}")

    log(f"Terminé avec {errors} erreur(s)." if errors else "Terminé.")
    return stats_list, errors


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args(argv: Optional[list[str]] = None):
    parser = argparse.ArgumentParser(description="Standardisation premium King Klown / UCKK pour .docx.")
    parser.add_argument("input", nargs="?", help="Fichier .docx ou dossier. Si absent : ouvre le GUI.")
    parser.add_argument("--output", "-o", default="standardises", help="Dossier de sortie si --in-place n'est pas utilisé")
    parser.add_argument("--recursive", "-r", action="store_true", help="Traiter aussi les sous-dossiers")
    parser.add_argument("--in-place", action="store_true", help="Écraser les originaux avec backup automatique")
    parser.add_argument("--suffix", default="_KK_PREMIUM", help="Suffixe ajouté aux fichiers générés")

    parser.add_argument("--page-size", choices=["letter", "a4"], default="letter")
    parser.add_argument("--margin-mode", choices=["centered", "left", "recto-verso"], default="centered")
    parser.add_argument("--header-layout", choices=["chapter", "brand-chapter", "centered", "split", "minimal", "none"], default="chapter", help="Header. Recommandé : chapter pour afficher le chapitre courant.")
    parser.add_argument("--body-alignment", choices=["left", "smart-justify", "justify"], default="left", help="Alignement du corps. Recommandé : left pour éviter les blancs étirés.")

    parser.add_argument("--toc", action="store_true", help="Insérer une table des matières")
    parser.add_argument("--no-header-footer", action="store_true", help="Ne pas ajouter header/footer")
    parser.add_argument("--no-chapter-breaks", action="store_true", help="Ne pas forcer les chapitres en nouvelle page")
    parser.add_argument("--keep-google-formatting", action="store_true", help="Conserver le formatage direct Google Docs")
    parser.add_argument("--no-tables", action="store_true", help="Ne pas standardiser les tableaux")
    parser.add_argument("--no-callout-boxes", action="store_true", help="Ne pas encadrer les blocs pédagogiques")
    parser.add_argument("--body-font", default="EB Garamond")
    parser.add_argument("--heading-font", default="EB Garamond")
    parser.add_argument("--footer-text", default="Univers-Cité King Klown")
    parser.add_argument("--dry-run", action="store_true", help="Analyser sans sauvegarder")
    parser.add_argument("--pdf", action="store_true", help="Exporter aussi PDF via LibreOffice")
    return parser.parse_args(argv)


def config_from_args(args) -> PremiumConfig:
    cfg = PremiumConfig()
    cfg.page_size = args.page_size
    cfg.margin_mode = args.margin_mode
    cfg.header_layout = args.header_layout
    cfg.body_alignment = args.body_alignment
    cfg.add_toc = args.toc
    cfg.add_header_footer = not args.no_header_footer
    cfg.chapter_page_breaks = not args.no_chapter_breaks
    cfg.clear_google_run_formatting = not args.keep_google_formatting
    cfg.standardize_tables = not args.no_tables
    cfg.add_callout_boxes = not args.no_callout_boxes
    cfg.body_font = args.body_font
    cfg.heading_font = args.heading_font
    cfg.footer_text = args.footer_text
    cfg.export_pdf = args.pdf
    return cfg


def cli_main(args) -> int:
    if not args.input:
        launch_gui()
        return 0

    input_root = Path(args.input).expanduser().resolve()
    output_root = Path(args.output).expanduser().resolve()
    cfg = config_from_args(args)

    _, errors = run_batch(
        input_root=input_root,
        output_root=output_root,
        cfg=cfg,
        recursive=args.recursive,
        in_place=args.in_place,
        suffix=args.suffix,
        dry_run=args.dry_run,
    )
    return 1 if errors else 0


# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------

class PremiumGUI(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("King Klown Premium Builder")
        self.geometry("880x640")
        self.minsize(780, 580)

        self.log_queue: queue.Queue[str] = queue.Queue()
        self.worker_thread: Optional[threading.Thread] = None

        self.input_path = tk.StringVar()
        self.output_path = tk.StringVar(value=str(Path.cwd() / "standardises"))

        self.suffix = tk.StringVar(value="_KK_PREMIUM")
        self.body_font = tk.StringVar(value="EB Garamond")
        self.heading_font = tk.StringVar(value="EB Garamond")
        self.footer_text = tk.StringVar(value="Univers-Cité King Klown")

        self.page_size = tk.StringVar(value="letter")
        self.margin_mode = tk.StringVar(value="centered")
        self.header_layout = tk.StringVar(value="chapter")
        self.body_alignment = tk.StringVar(value="left")

        self.recursive = tk.BooleanVar(value=False)
        self.in_place = tk.BooleanVar(value=False)
        self.add_toc = tk.BooleanVar(value=True)
        self.header_footer = tk.BooleanVar(value=True)
        self.chapter_breaks = tk.BooleanVar(value=True)
        self.clear_google_formatting = tk.BooleanVar(value=True)
        self.standardize_tables_var = tk.BooleanVar(value=True)
        self.callout_boxes = tk.BooleanVar(value=True)
        self.dry_run = tk.BooleanVar(value=False)
        self.export_pdf_var = tk.BooleanVar(value=False)

        self._build_ui()
        self.after(120, self._pump_log_queue)

    def _build_ui(self) -> None:
        root = ttk.Frame(self, padding=12)
        root.pack(fill="both", expand=True)

        source = ttk.LabelFrame(root, text="Source")
        source.pack(fill="x", pady=(0, 10))
        ttk.Label(source, text="Fichier ou dossier .docx").grid(row=0, column=0, sticky="w", padx=8, pady=8)
        ttk.Entry(source, textvariable=self.input_path).grid(row=0, column=1, sticky="ew", padx=8, pady=8)
        ttk.Button(source, text="Fichier…", command=self._choose_file).grid(row=0, column=2, padx=4, pady=8)
        ttk.Button(source, text="Dossier…", command=self._choose_folder).grid(row=0, column=3, padx=8, pady=8)
        source.columnconfigure(1, weight=1)

        out = ttk.LabelFrame(root, text="Sortie")
        out.pack(fill="x", pady=(0, 10))
        ttk.Label(out, text="Dossier de sortie").grid(row=0, column=0, sticky="w", padx=8, pady=8)
        self.output_entry = ttk.Entry(out, textvariable=self.output_path)
        self.output_entry.grid(row=0, column=1, sticky="ew", padx=8, pady=8)
        ttk.Button(out, text="Choisir…", command=self._choose_output).grid(row=0, column=2, padx=8, pady=8)
        ttk.Label(out, text="Suffixe").grid(row=1, column=0, sticky="w", padx=8, pady=8)
        ttk.Entry(out, textvariable=self.suffix, width=18).grid(row=1, column=1, sticky="w", padx=8, pady=8)
        out.columnconfigure(1, weight=1)

        layout = ttk.LabelFrame(root, text="Layout premium")
        layout.pack(fill="x", pady=(0, 10))

        ttk.Label(layout, text="Page").grid(row=0, column=0, sticky="w", padx=8, pady=6)
        ttk.Combobox(layout, textvariable=self.page_size, values=["letter", "a4"], state="readonly", width=14).grid(row=0, column=1, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Marges").grid(row=0, column=2, sticky="w", padx=8, pady=6)
        ttk.Combobox(
            layout,
            textvariable=self.margin_mode,
            values=["centered", "left", "recto-verso"],
            state="readonly",
            width=16
        ).grid(row=0, column=3, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Header").grid(row=0, column=4, sticky="w", padx=8, pady=6)
        ttk.Combobox(
            layout,
            textvariable=self.header_layout,
            values=["chapter", "brand-chapter", "centered", "split", "minimal", "none"],
            state="readonly",
            width=16
        ).grid(row=0, column=5, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Alignement").grid(row=1, column=4, sticky="w", padx=8, pady=6)
        ttk.Combobox(
            layout,
            textvariable=self.body_alignment,
            values=["left", "smart-justify", "justify"],
            state="readonly",
            width=14
        ).grid(row=1, column=5, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Police corps").grid(row=1, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(layout, textvariable=self.body_font, width=18).grid(row=1, column=1, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Police titres").grid(row=1, column=2, sticky="w", padx=8, pady=6)
        ttk.Entry(layout, textvariable=self.heading_font, width=18).grid(row=1, column=3, sticky="w", padx=8, pady=6)

        ttk.Label(layout, text="Pied de page").grid(row=2, column=0, sticky="w", padx=8, pady=6)
        ttk.Entry(layout, textvariable=self.footer_text).grid(row=2, column=1, columnspan=5, sticky="ew", padx=8, pady=6)
        layout.columnconfigure(5, weight=1)

        checks = ttk.LabelFrame(root, text="Traitement")
        checks.pack(fill="x", pady=(0, 10))

        checkbox_data = [
            ("Récursif", self.recursive),
            ("Écraser originaux avec backup", self.in_place),
            ("Table des matières", self.add_toc),
            ("En-tête / pied de page", self.header_footer),
            ("Sauts de page chapitres", self.chapter_breaks),
            ("Nettoyer formatage Google Docs", self.clear_google_formatting),
            ("Standardiser tableaux", self.standardize_tables_var),
            ("Encadrés pédagogiques", self.callout_boxes),
            ("Dry-run seulement", self.dry_run),
            ("Exporter PDF via LibreOffice", self.export_pdf_var),
        ]

        for idx, (label, var) in enumerate(checkbox_data):
            r = idx // 2
            c = (idx % 2) * 2
            ttk.Checkbutton(checks, text=label, variable=var, command=self._toggle_in_place).grid(
                row=r, column=c, columnspan=2, sticky="w", padx=8, pady=4
            )

        actions = ttk.Frame(root)
        actions.pack(fill="x", pady=(0, 10))
        self.run_button = ttk.Button(actions, text="Standardiser premium", command=self._start_processing)
        self.run_button.pack(side="left")
        ttk.Button(actions, text="Effacer le log", command=self._clear_log).pack(side="left", padx=8)
        self.progress = ttk.Progressbar(actions, mode="indeterminate")
        self.progress.pack(side="right", fill="x", expand=True, padx=(8, 0))

        log_frame = ttk.LabelFrame(root, text="Log")
        log_frame.pack(fill="both", expand=True)
        self.log_text = tk.Text(log_frame, wrap="word", height=14)
        self.log_text.pack(side="left", fill="both", expand=True)
        scroll = ttk.Scrollbar(log_frame, command=self.log_text.yview)
        scroll.pack(side="right", fill="y")
        self.log_text.configure(yscrollcommand=scroll.set)

    def _choose_file(self) -> None:
        path = filedialog.askopenfilename(
            title="Choisir un fichier .docx",
            filetypes=[("Word documents", "*.docx"), ("Tous les fichiers", "*.*")]
        )
        if path:
            self.input_path.set(path)

    def _choose_folder(self) -> None:
        path = filedialog.askdirectory(title="Choisir un dossier")
        if path:
            self.input_path.set(path)

    def _choose_output(self) -> None:
        path = filedialog.askdirectory(title="Choisir le dossier de sortie")
        if path:
            self.output_path.set(path)

    def _toggle_in_place(self) -> None:
        self.output_entry.configure(state="disabled" if self.in_place.get() else "normal")

    def _clear_log(self) -> None:
        self.log_text.delete("1.0", "end")

    def _log(self, msg: str) -> None:
        self.log_text.insert("end", msg + "\n")
        self.log_text.see("end")

    def _thread_log(self, msg: str) -> None:
        self.log_queue.put(msg)

    def _pump_log_queue(self) -> None:
        try:
            while True:
                msg = self.log_queue.get_nowait()
                if msg == "__DONE__":
                    self.progress.stop()
                    self.run_button.configure(state="normal")
                else:
                    self._log(msg)
        except queue.Empty:
            pass
        self.after(120, self._pump_log_queue)

    def _start_processing(self) -> None:
        if self.worker_thread and self.worker_thread.is_alive():
            messagebox.showinfo("Traitement en cours", "Un traitement est déjà en cours.")
            return

        source = self.input_path.get().strip()
        if not source:
            messagebox.showerror("Source manquante", "Choisis un fichier .docx ou un dossier.")
            return

        input_root = Path(source).expanduser().resolve()
        output_root = Path(self.output_path.get().strip() or "standardises").expanduser().resolve()

        if self.in_place.get():
            confirm = messagebox.askyesno(
                "Confirmer le mode in-place",
                "Le mode in-place écrase les fichiers originaux, avec une sauvegarde .bak automatique.\n\nContinuer ?"
            )
            if not confirm:
                return

        cfg = PremiumConfig(
            page_size=self.page_size.get(),
            margin_mode=self.margin_mode.get(),
            header_layout=self.header_layout.get(),
            body_alignment=self.body_alignment.get(),
            body_font=self.body_font.get().strip() or "EB Garamond",
            heading_font=self.heading_font.get().strip() or "EB Garamond",
            footer_text=self.footer_text.get().strip() or "Univers-Cité King Klown",
            add_toc=self.add_toc.get(),
            add_header_footer=self.header_footer.get(),
            chapter_page_breaks=self.chapter_breaks.get(),
            clear_google_run_formatting=self.clear_google_formatting.get(),
            standardize_tables=self.standardize_tables_var.get(),
            add_callout_boxes=self.callout_boxes.get(),
            export_pdf=self.export_pdf_var.get(),
        )

        self.run_button.configure(state="disabled")
        self.progress.start(10)
        self._log("—" * 72)
        self._log(
            f"Premium | page={cfg.page_size} | marges={cfg.margin_mode} | "
            f"header={cfg.header_layout} | alignement={cfg.body_alignment} | brand=#1E6864 | corps={cfg.body_font} | titres={cfg.heading_font}"
        )

        def worker() -> None:
            try:
                run_batch(
                    input_root=input_root,
                    output_root=output_root,
                    cfg=cfg,
                    recursive=self.recursive.get(),
                    in_place=self.in_place.get(),
                    suffix=self.suffix.get().strip() or "_KK_PREMIUM",
                    dry_run=self.dry_run.get(),
                    logger=self._thread_log,
                )
            except Exception as exc:
                self._thread_log(f"ERREUR GÉNÉRALE : {exc}")
            finally:
                self._thread_log("__DONE__")

        self.worker_thread = threading.Thread(target=worker, daemon=True)
        self.worker_thread.start()


def launch_gui() -> None:
    app = PremiumGUI()
    app.mainloop()


def main(argv: Optional[list[str]] = None) -> int:
    args = parse_args(argv)
    return cli_main(args)


if __name__ == "__main__":
    raise SystemExit(main())
